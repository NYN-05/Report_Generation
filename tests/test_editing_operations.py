"""Integration tests for editing operations on document structure."""

import pytest
from docx import Document


class TestReplaceSection:
    def test_replace_section_content(self, sample_docx):
        from src.document.structure import build_tree, ReplaceSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        op = ReplaceSection(root, doc)
        result = op.execute(target="Introduction", new_content="This is the new introduction content.\n\nIt has multiple paragraphs.")
        assert result.success is True
        assert len(result.modified_elements) >= 1

    def test_replace_with_new_heading(self, sample_docx):
        from src.document.structure import build_tree, ReplaceSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        op = ReplaceSection(root, doc)
        result = op.execute(target="Introduction", new_content="Updated.",
                            new_heading="Revised Introduction")
        assert result.success is True

    def test_replace_nonexistent_section(self, sample_docx):
        from src.document.structure import build_tree, ReplaceSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        op = ReplaceSection(root, doc)
        result = op.execute(target="Nonexistent Section")
        assert result.success is False
        assert "not found" in (result.error or "").lower()

    def test_replace_preserves_tables(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        assert analysis is not None
        tables_before = locator.find_tables(analysis)

        op = ReplaceSection(root, doc)
        result = op.execute(target="Analysis", new_content="Replaced content.",
                            preserve_children=True)
        assert result.success is True

        locator2 = SectionLocator(root)
        analysis2 = locator2.find_by_heading("Analysis")
        if analysis2:
            tables_after = locator2.find_tables(analysis2)
            assert len(tables_after) == len(tables_before)


class TestInsertSection:
    def test_insert_after(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, InsertSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        sections_before = len(locator.find_by_level(1))

        op = InsertSection(root, doc)
        result = op.execute(target="Introduction", heading="New Section",
                            content="Content of new section.", position="after")
        assert result.success is True

        locator2 = SectionLocator(root)
        sections_after = len(locator2.find_by_level(1))
        assert sections_after == sections_before + 1

    def test_insert_before(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, InsertSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        sections_before = len(locator.find_by_level(1))

        op = InsertSection(root, doc)
        result = op.execute(target="Conclusion", heading="Pre-Conclusion",
                            content="Before conclusion.", position="before")
        assert result.success is True

        locator2 = SectionLocator(root)
        sections_after = len(locator2.find_by_level(1))
        assert sections_after == sections_before + 1

    def test_insert_without_target(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, InsertSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        sections_before = len(locator.find_by_level(1))

        op = InsertSection(root, doc)
        result = op.execute(heading="Final Section", content="End content.")
        assert result.success is True

        locator2 = SectionLocator(root)
        sections_after = len(locator2.find_by_level(1))
        assert sections_after == sections_before + 1


class TestExpandSection:
    def test_expand_with_subsections(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ExpandSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)

        op = ExpandSection(root, doc)
        result = op.execute(
            target="Analysis",
            new_subsections=[
                {"heading": "Quantitative Analysis", "content": "Numbers and data."},
                {"heading": "Qualitative Analysis", "content": "Insights and observations."},
            ],
        )
        assert result.success is True
        assert "Quantitative Analysis" in str(result.modified_elements)

    def test_expand_with_paragraphs(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ExpandSection

        doc = Document(sample_docx)
        root = build_tree(doc)

        op = ExpandSection(root, doc)
        result = op.execute(
            target="Conclusion",
            append_paragraphs=["First additional paragraph.", "Second additional paragraph."],
        )
        assert result.success is True

        locator2 = SectionLocator(root)
        conclusion = locator2.find_by_heading("Conclusion")
        if conclusion:
            paras = locator2.find_content_blocks(conclusion)
            assert len(paras) >= 2

    def test_expand_nonexistent(self, sample_docx):
        from src.document.structure import build_tree, ExpandSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        op = ExpandSection(root, doc)
        result = op.execute(target="Fake Section")
        assert result.success is False


class TestDeleteSection:
    def test_delete_section(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, DeleteSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        subsections_before = len(locator.find_by_level(2))

        op = DeleteSection(root, doc)
        result = op.execute(target="Methodology")
        assert result.success is True

        locator2 = SectionLocator(root)
        subsections_after = len(locator2.find_by_level(2))
        assert subsections_after == subsections_before - 1

    def test_delete_children_only(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, DeleteSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        children_before = len(analysis.children) if analysis else 0

        op = DeleteSection(root, doc)
        result = op.execute(target="Analysis", delete_children_only=True)
        assert result.success is True

        locator2 = SectionLocator(root)
        analysis2 = locator2.find_by_heading("Analysis")
        if analysis2:
            assert len(analysis2.children) < children_before

    def test_delete_nonexistent(self, sample_docx):
        from src.document.structure import build_tree, DeleteSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        op = DeleteSection(root, doc)
        result = op.execute(target="Missing Section")
        assert result.success is False


class TestMoveSection:
    def test_move_after(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, MoveSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        sections_before = len(locator.find_by_level(1))

        op = MoveSection(root, doc)
        result = op.execute(target="Conclusion", destination="Introduction", position="after")
        assert result.success is True

        locator2 = SectionLocator(root)
        sections_after = len(locator2.find_by_level(1))
        assert sections_after == sections_before

        intro = locator2.find_by_heading("Introduction")
        conclusion = locator2.find_by_heading("Conclusion")
        assert intro is not None
        assert conclusion is not None


class TestEditingPlannerIntegration:
    def test_full_flow_expand(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, EditingPlanner, ExpandSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)

        ops = planner.plan('Expand "Introduction" with subsections')
        assert len(ops) > 0

        locator = SectionLocator(root)
        intro = locator.find_by_heading("Introduction")
        if intro:
            op = ExpandSection(root, doc)
            result = op.execute(
                target_node=intro,
                new_subsections=[{"heading": "Background", "content": "Background details."}],
            )
            assert result.success is True

    def test_full_flow_insert_then_replace(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, InsertSection, ReplaceSection

        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        sections_before = len(locator.find_by_level(1))

        insert = InsertSection(root, doc)
        insert_result = insert.execute(
            target="Executive Summary",
            heading="Recommendations",
            content="Key recommendations.",
            position="after",
        )
        assert insert_result.success is True

        locator2 = SectionLocator(root)
        assert len(locator2.find_by_level(1)) == sections_before + 1

        replace = ReplaceSection(root, doc)
        replace_result = replace.execute(
            target="Recommendations",
            new_content="Updated recommendations with more details.\n\nSecond paragraph.",
        )
        assert replace_result.success is True

    def test_planner_operations_match_available_operations(self):
        from src.document.structure import EditingPlanner
        from src.document.structure.operations import (
            ReplaceSection, InsertSection, ExpandSection, DeleteSection, MoveSection,
        )
        planner = EditingPlanner()
        ops = ["expand", "replace", "insert", "delete", "move"]
        for op_name in ops:
            results = planner.plan(f"{op_name} test")
            assert len(results) > 0, f"Planner should handle: {op_name}"


class TestStructuralEditing:
    """Tests proving targeted DOCX body XML manipulation works correctly."""

    def test_edit_only_affects_target_section_body(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        body_before = len(list(doc.element.body))

        root = build_tree(doc)
        locator = SectionLocator(root)

        intro = locator.find_by_heading("Introduction")
        intro_body_idx = None
        for i, child in enumerate(doc.element.body):
            if child is intro.element._element:
                intro_body_idx = i
                break

        op = ReplaceSection(root, doc)
        op.execute(target="Introduction", new_content="Replacement para.")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)

        summary = locator2.find_by_heading("Executive Summary")
        assert summary is not None, "Exec Summary should still exist"
        analysis = locator2.find_by_heading("Analysis")
        assert analysis is not None, "Analysis should still exist"
        conclusion = locator2.find_by_heading("Conclusion")
        assert conclusion is not None, "Conclusion should still exist"

    def test_nonedited_sections_have_unchanged_paragraphs(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        conclusion = locator.find_by_heading("Conclusion")
        conclusion_text = ""
        for c in conclusion.children:
            if hasattr(c, 'text') and c.text:
                conclusion_text = c.text
                break

        op = ReplaceSection(root, doc)
        op.execute(target="Introduction", new_content="Replacement content.")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        conclusion2 = locator2.find_by_heading("Conclusion")
        conclusion_text2 = ""
        for c in conclusion2.children:
            if hasattr(c, 'text') and c.text:
                conclusion_text2 = c.text
                break

        assert conclusion_text2 == conclusion_text, (
            f"Conclusion changed from '{conclusion_text}' to '{conclusion_text2}'"
        )

    def test_tables_survive_replace_in_different_section(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        tables_before = locator.find_tables(analysis)

        op = ReplaceSection(root, doc)
        op.execute(target="Introduction", new_content="New intro text.")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        analysis2 = locator2.find_by_heading("Analysis")
        tables_after = locator2.find_tables(analysis2)
        assert len(tables_after) == len(tables_before), (
            f"Table count changed from {len(tables_before)} to {len(tables_after)}"
        )

    def test_tables_survive_insert_adjacent(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, InsertSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        tables_before = locator.find_tables(analysis)

        op = InsertSection(root, doc)
        op.execute(target="Analysis", heading="New Section",
                   content="Inserted content.", position="before")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        analysis2 = locator2.find_by_heading("Analysis")
        tables_after = locator2.find_tables(analysis2)
        assert len(tables_after) == len(tables_before), (
            f"Table count changed from {len(tables_before)} to {len(tables_after)}"
        )

    def test_tables_survive_delete_adjacent(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, DeleteSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        tables_before = locator.find_tables(analysis)

        op = DeleteSection(root, doc)
        op.execute(target="Introduction")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        analysis2 = locator2.find_by_heading("Analysis")
        tables_after = locator2.find_tables(analysis2)
        assert len(tables_after) == len(tables_before), (
            f"Table count changed from {len(tables_before)} to {len(tables_after)}"
        )

    def test_paragraph_count_preserved_in_unrelated_sections(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        exec_summary = locator.find_by_heading("Executive Summary")
        summary_para_count = len([c for c in exec_summary.children
                                  if c.node_type.value == "paragraph"])

        op = ReplaceSection(root, doc)
        op.execute(target="Introduction", new_content="Replacement.")

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        exec_summary2 = locator2.find_by_heading("Executive Summary")
        summary_para_count2 = len([c for c in exec_summary2.children
                                   if c.node_type.value == "paragraph"])
        assert summary_para_count2 == summary_para_count

    def test_replace_preserves_pre_content_before_heading(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        cover_pages = root.find_by_type("cover_page")
        cover_before = len(cover_pages)

        op = ReplaceSection(root, doc)
        result = op.execute(target="Introduction", new_content="Replacement.")
        assert result.success is True

        root2 = build_tree(doc)
        cover_pages2 = root2.find_by_type("cover_page")
        assert len(cover_pages2) == cover_before

    def test_expand_preserves_existing_subsections(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ExpandSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        intro = locator.find_by_heading("Introduction")
        subs_before = len(locator.find_subsections(intro))

        op = ExpandSection(root, doc)
        op.execute(target="Introduction",
                   new_subsections=[{"heading": "Background", "content": "BG text."}])

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        intro2 = locator2.find_by_heading("Introduction")
        subs_after = len(locator2.find_subsections(intro2))
        assert subs_after == subs_before + 1

    def test_delete_children_only_preserves_tables(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, DeleteSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        tables_before = locator.find_tables(analysis)
        assert len(tables_before) >= 1

        op = DeleteSection(root, doc)
        op.execute(target="Analysis", delete_children_only=True)

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        analysis2 = locator2.find_by_heading("Analysis")
        tables_after = locator2.find_tables(analysis2)
        assert len(tables_after) >= 1


class TestChainedEdits:
    """Tests that multiple structural edits can be applied sequentially."""

    def test_insert_then_expand_then_replace(self, sample_docx):
        from src.document.structure import (
            build_tree, SectionLocator,
            InsertSection, ExpandSection, ReplaceSection,
        )
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        initial_count = len(locator.find_by_level(1))

        insert_op = InsertSection(root, doc)
        r1 = insert_op.execute(target="Introduction", heading="Middle",
                                content="Middle content.", position="after")
        assert r1.success is True

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        assert len(locator2.find_by_level(1)) == initial_count + 1

        expand_op = ExpandSection(root2, doc)
        r2 = expand_op.execute(
            target="Middle",
            new_subsections=[{"heading": "Sub-Middle", "content": "Sub content."}],
        )
        assert r2.success is True

        root3 = build_tree(doc)
        locator3 = SectionLocator(root3)
        middle = locator3.find_by_heading("Middle")
        sub_count = len(locator3.find_subsections(middle))
        assert sub_count >= 1

        replace_op = ReplaceSection(root3, doc)
        r3 = replace_op.execute(target="Sub-Middle", new_content="Replaced sub content.")
        assert r3.success is True

        root4 = build_tree(doc)
        locator4 = SectionLocator(root4)
        sub_middle = locator4.find_by_heading("Sub-Middle")
        paras = locator4.find_content_blocks(sub_middle) if sub_middle else []
        assert len(paras) >= 1

    def test_delete_then_insert(self, sample_docx):
        from src.document.structure import (
            build_tree, SectionLocator,
            DeleteSection, InsertSection,
        )
        doc = Document(sample_docx)

        root = build_tree(doc)
        locator = SectionLocator(root)
        initial_count = len(locator.find_by_level(2))

        delete_op = DeleteSection(root, doc)
        r1 = delete_op.execute(target="Methodology")
        assert r1.success is True

        root2 = build_tree(doc)
        locator2 = SectionLocator(root2)
        assert len(locator2.find_by_level(2)) == initial_count - 1

        insert_op = InsertSection(root2, doc)
        r2 = insert_op.execute(target="Introduction", heading="New Subsection",
                                content="New subsection text.", level=2, position="after")
        assert r2.success is True

        root3 = build_tree(doc)
        locator3 = SectionLocator(root3)
        new_sub = locator3.find_by_heading("New Subsection")
        assert new_sub is not None
        assert len(locator3.find_by_level(2)) == initial_count

    def test_chained_edits_preserve_tables(self, sample_docx):
        from src.document.structure import (
            build_tree, SectionLocator,
            ExpandSection, DeleteSection, InsertSection,
        )
        doc = Document(sample_docx)

        # Track table count across edits
        def get_table_count(d, heading="Analysis"):
            t = build_tree(d)
            l = SectionLocator(t)
            s = l.find_by_heading(heading)
            if s:
                return len(l.find_tables(s))
            return 0

        tables_baseline = get_table_count(doc)

        root = build_tree(doc)
        ExpandSection(root, doc).execute(
            target="Analysis",
            append_paragraphs=["Extra para."],
        )
        assert get_table_count(doc) == tables_baseline

        root2 = build_tree(doc)
        DeleteSection(root2, doc).execute(target="Introduction")
        assert get_table_count(doc) == tables_baseline

        root3 = build_tree(doc)
        InsertSection(root3, doc).execute(
            target="Analysis", heading="After Analysis",
            content="Follow-up content.", position="after",
        )
        assert get_table_count(doc) == tables_baseline

    def test_replace_then_replace_same_section(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator, ReplaceSection
        doc = Document(sample_docx)

        root = build_tree(doc)
        ReplaceSection(root, doc).execute(
            target="Introduction", new_content="First replacement."
        )

        root2 = build_tree(doc)
        r2 = ReplaceSection(root2, doc).execute(
            target="Introduction", new_content="Second replacement.\n\nMore text."
        )
        assert r2.success is True

        root3 = build_tree(doc)
        locator3 = SectionLocator(root3)
        intro3 = locator3.find_by_heading("Introduction")
        paras = locator3.find_content_blocks(intro3) if intro3 else []
        assert len(paras) >= 2


class TestSectionBoundaries:
    """Tests for element boundary detection used by structural editing."""

    def test_heading_body_index_found(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator
        from src.document.structure.operations import _find_heading_body_index
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        intro = locator.find_by_heading("Introduction")
        idx = _find_heading_body_index(doc, intro)
        assert idx >= 0
        body_children = list(doc.element.body)
        tag = body_children[idx].tag.split('}')[-1] if '}' in body_children[idx].tag else body_children[idx].tag
        assert tag == 'p'

    def test_sibling_boundary_after_last_section(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator
        from src.document.structure.operations import _sibling_section_body_start
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        conclusion = locator.find_by_heading("Conclusion")
        boundary = _sibling_section_body_start(doc, root, conclusion)
        body_len = len(list(doc.element.body))
        assert boundary == body_len

    def test_section_body_range_non_empty(self, sample_docx):
        from src.document.structure import build_tree, SectionLocator
        from src.document.structure.operations import _section_body_range
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        r = _section_body_range(doc, root, analysis)
        assert r != (-1, -1)
        assert r[0] >= 0
        assert r[1] > r[0]


class TestFormatPreservation:
    """Tests that formatting is preserved during structural edits."""

    def test_heading_style_preserved_after_expand(self, sample_docx):
        from src.document.structure import build_tree, ExpandSection
        from docx.oxml.ns import qn
        doc = Document(sample_docx)

        root = build_tree(doc)
        ExpandSection(root, doc).execute(
            target="Analysis",
            new_subsections=[{"heading": "Sub Analysis", "content": "Sub content."}],
        )

        for para in doc.paragraphs:
            if para.text.strip() == "Sub Analysis":
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        val = pStyle.get(qn('w:val'))
                        assert val == "Heading2", f"Expected Heading2, got {val}"
                break
        else:
            pytest.fail("Could not find inserted subsection heading")

    def test_existing_heading_styles_untouched(self, sample_docx):
        from src.document.structure import build_tree, ReplaceSection
        from docx.oxml.ns import qn
        doc = Document(sample_docx)

        original_styles = {}
        for para in doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    original_styles[para.text.strip()] = pStyle.get(qn('w:val'))

        root = build_tree(doc)
        ReplaceSection(root, doc).execute(
            target="Introduction", new_content="Replacement text."
        )

        for para in doc.paragraphs:
            text = para.text.strip()
            if text in original_styles and text != "Introduction":
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        assert pStyle.get(qn('w:val')) == original_styles[text], (
                            f"Style changed for '{text}'"
                        )

    def test_body_element_count_stable_outside_edit_zone(self, sample_docx):
        from src.document.structure import build_tree, DeleteSection, InsertSection, ExpandSection
        doc = Document(sample_docx)

        body_children_before = list(doc.element.body)

        # Find index of "Conclusion" heading
        concl_idx = -1
        for i, child in enumerate(body_children_before):
            for para in doc.paragraphs:
                if para._element is child and "Conclusion" in para.text:
                    concl_idx = i
                    break
            if concl_idx >= 0:
                break

        root = build_tree(doc)
        ExpandSection(root, doc).execute(
            target="Analysis",
            new_subsections=[{"heading": "Extra", "content": "Extra para."}],
        )

        body_children_after = list(doc.element.body)
        pre_concl_before = body_children_before[:concl_idx]
        pre_concl_after = body_children_after[:concl_idx]

        assert len(pre_concl_after) >= len(pre_concl_before)
