"""Unit tests for Document Structure Model, Section Locator, and Editing Planner."""

import pytest
from docx import Document


class TestBuildTree:
    def test_build_tree_creates_document_node(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        assert root is not None
        assert root.node_type.value == "document"

    def test_build_tree_detects_sections(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        sections = root.find_by_type("section")
        assert len(sections) >= 4

    def test_build_tree_detects_cover_page(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        cover_pages = root.find_by_type("cover_page")
        assert len(cover_pages) > 0

    def test_build_tree_detects_toc(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        tocs = root.find_by_type("toc")
        assert len(tocs) > 0

    def test_section_has_heading(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        sections = root.find_by_type("section")
        headings = [s.metadata.get("heading") for s in sections]
        assert "Executive Summary" in headings
        assert "Introduction" in headings
        assert "Analysis" in headings
        assert "Conclusion" in headings

    def test_section_levels(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        sections = root.find_by_type("section")
        l1 = [s for s in sections if s.metadata.get("level") == 1]
        l2 = [s for s in sections if s.metadata.get("level") == 2]
        assert len(l1) >= 4
        assert len(l2) >= 1

    def test_paragraphs_inside_sections(self, sample_docx, build_tree):
        doc = Document(sample_docx)
        root = build_tree(doc)
        sections = root.find_by_type("section")
        paras_found = False
        for sec in sections:
            for child in sec.children:
                if child.node_type.value == "paragraph":
                    paras_found = True
                    break
        assert paras_found


class TestSectionLocator:
    def test_find_by_heading_exact(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        section = locator.find_by_heading("Analysis")
        assert section is not None
        assert section.metadata.get("heading") == "Analysis"

    def test_find_by_heading_fuzzy(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        results = locator.find_by_heading_fuzzy("Intro")
        assert len(results) >= 1
        assert "Introduction" in [s.metadata.get("heading") for s in results]

    def test_find_by_level(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        l1 = locator.find_by_level(1)
        assert len(l1) >= 4

    def test_find_content_blocks(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        section = locator.find_by_heading("Introduction")
        assert section is not None
        paras = locator.find_content_blocks(section)
        assert len(paras) >= 1

    def test_get_hierarchy(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        hierarchy = locator.get_hierarchy()
        assert len(hierarchy) >= 4
        assert all("heading" in h for h in hierarchy)
        assert all("level" in h for h in hierarchy)

    def test_get_path(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        section = locator.find_by_heading("Methodology")
        if section:
            path = locator.get_path(section)
            assert "Methodology" in path

    def test_find_subsections(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        intro = locator.find_by_heading("Introduction")
        if intro:
            subs = locator.find_subsections(intro)
            assert len(subs) >= 1

    def test_find_tables(self, sample_docx, build_tree, SectionLocator):
        doc = Document(sample_docx)
        root = build_tree(doc)
        locator = SectionLocator(root)
        analysis = locator.find_by_heading("Analysis")
        if analysis:
            tables = locator.find_tables(analysis)
            assert len(tables) >= 1


class TestEditingPlanner:
    def test_plan_expand(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Expand "Analysis"')
        assert len(ops) >= 1
        assert ops[0].operation == "expand"
        assert "analysis" in ops[0].target.lower()

    def test_plan_replace(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Replace "Introduction" with new content')
        assert len(ops) >= 1
        assert ops[0].operation == "replace"

    def test_plan_insert(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Insert "Recommendations" after "Analysis"')
        assert len(ops) >= 1
        assert ops[0].operation == "insert"

    def test_plan_delete(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Delete "Methodology"')
        assert len(ops) >= 1
        assert ops[0].operation == "delete"

    def test_plan_move(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Move "Conclusion" after "Introduction"')
        assert len(ops) >= 1
        assert ops[0].operation == "move"

    def test_explain_plan(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan('Expand "Analysis"')
        explanation = planner.explain_plan(ops)
        assert "EXPAND" in explanation.upper()
        assert "analysis" in explanation.lower()

    def test_empty_instruction(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan("")
        assert len(ops) == 0

    def test_unrecognized_instruction(self, sample_docx, build_tree, EditingPlanner):
        doc = Document(sample_docx)
        root = build_tree(doc)
        planner = EditingPlanner(root)
        ops = planner.plan("do something completely random and weird 42")
        assert len(ops) == 0


class TestNodeModel:
    def test_section_node(self):
        from src.document.structure import SectionNode
        node = SectionNode(heading="Test", level=2)
        assert node.heading == "Test"
        assert node.level == 2
        assert node.node_type.value == "section"

    def test_paragraph_node(self):
        from src.document.structure import ParagraphNode
        node = ParagraphNode(text="Hello world")
        assert node.text == "Hello world"
        assert node.node_type.value == "paragraph"

    def test_node_hierarchy(self):
        from src.document.structure import DocumentNode, SectionNode, ParagraphNode
        root = DocumentNode()
        sec = SectionNode(heading="Chapter 1")
        para = ParagraphNode(text="Content")
        sec.add_child(para)
        root.add_child(sec)
        assert para.parent is sec
        assert sec.parent is root
        found = root.find_by_id(sec.node_id)
        assert found is sec

    def test_find_by_type(self):
        from src.document.structure import DocumentNode, SectionNode, ParagraphNode
        root = DocumentNode()
        s1 = SectionNode(heading="A")
        s2 = SectionNode(heading="B")
        p1 = ParagraphNode(text="x")
        s1.add_child(p1)
        root.add_child(s1)
        root.add_child(s2)
        all_sections = root.find_by_type("section")
        assert len(all_sections) == 2
        all_paras = root.find_by_type("paragraph")
        assert len(all_paras) == 1

    def test_node_to_dict(self):
        from src.document.structure import SectionNode
        node = SectionNode(heading="Test")
        d = node.to_dict()
        assert d["node_type"] == "section"
        assert d["metadata"]["heading"] == "Test"
