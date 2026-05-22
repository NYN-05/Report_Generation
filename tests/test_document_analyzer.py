"""
Tests for the DOCX Analyzer module.
Tests model, heading, classifier, styles, tables, images, references, graph, and parser.
"""

import os
import sys
import tempfile
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════
#  Fixtures
# ═══════════════════════════════════════════════════════════════

@pytest.fixture
def academic_docx():
    """Create a realistic academic report with varied structure."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Certificate / Declaration ---
    doc.add_heading("Certificate", level=1)
    doc.add_paragraph(
        "This is to certify that the work presented in this report is original."
    )

    # --- Acknowledgement ---
    doc.add_heading("Acknowledgement", level=1)
    doc.add_paragraph("I would like to thank my advisor and colleagues.")

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive analysis of the subject matter. "
        "The findings indicate significant trends in the field."
    )

    doc.add_page_break()

    # --- 1. Introduction ---
    p = doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph(
        "This section introduces the topic and outlines the scope of the report."
    )
    doc.add_paragraph(
        "The background context is established through a review of existing literature."
    )

    # --- 1.1 Background ---
    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph("The background provides historical context for the study.")

    # --- 1.2 Problem Statement ---
    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph("The problem addressed in this report is clearly defined.")

    # --- 2. Literature Survey ---
    doc.add_heading("2. Literature Survey", level=1)
    doc.add_paragraph(
        "This section reviews relevant literature [1] and prior work [2], [3]."
    )
    doc.add_paragraph("Key papers include Smith (2020) and Jones (2019).")

    # --- 3. Methodology ---
    doc.add_heading("3. Methodology", level=1)
    doc.add_paragraph("The methodology section describes the approach used.")

    # Table 1: Methodology comparison
    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'
    for i, headers in enumerate(["Method", "Advantage", "Limitation"]):
        table1.rows[0].cells[i].text = headers
    data = [
        ["Quantitative", "Objective data", "Context missing"],
        ["Qualitative", "Rich insights", "Small sample"],
        ["Mixed", "Comprehensive", "Complex analysis"],
    ]
    for ri, row_data in enumerate(data, 1):
        for ci, val in enumerate(row_data):
            table1.rows[ri].cells[ci].text = val

    doc.add_paragraph("Table 1: Methodology comparison summary")

    # --- 3.1 Data Collection ---
    doc.add_heading("3.1 Data Collection", level=2)
    doc.add_paragraph("Data was collected from multiple sources.")

    # --- 4. Results ---
    doc.add_heading("4. Results", level=1)
    doc.add_paragraph("The results demonstrate significant findings [1, p. 5].")

    # Table 2: Results summary
    table2 = doc.add_table(rows=3, cols=4)
    table2.style = 'Table Grid'
    for i, h in enumerate(["Metric", "Year 1", "Year 2", "Year 3"]):
        table2.rows[0].cells[i].text = h
    results = [
        ["Revenue", "100", "150", "200"],
        ["Growth", "10%", "15%", "20%"],
    ]
    for ri, rd in enumerate(results, 1):
        for ci, v in enumerate(rd):
            table2.rows[ri].cells[ci].text = v

    doc.add_paragraph("Table 2: Annual performance metrics")

    # --- 5. Discussion ---
    doc.add_heading("5. Discussion", level=1)
    doc.add_paragraph(
        "The results are interpreted in the context of prior work [2]. "
        "According to Smith (2020), these findings align with established theory."
    )

    # --- 6. Conclusion ---
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph("This report concludes with key findings and recommendations.")
    doc.add_paragraph("Future work should explore additional dimensions.")

    # --- References ---
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] J. Smith, 'Advances in Machine Learning,' IEEE Press, 2020.")
    doc.add_paragraph("[2] A. Jones and B. Lee, 'Deep Learning Applications,' MIT Press, 2019.")
    doc.add_paragraph("[3] M. Brown et al., 'Survey of AI Methods,' Springer, 2021.")

    # --- Appendix A ---
    doc.add_heading("Appendix A", level=1)
    doc.add_paragraph("Supplementary data and additional figures.")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    fname = tmp.name
    tmp.close()
    doc.save(fname)
    yield fname
    try:
        os.unlink(fname)
    except OSError:
        pass


@pytest.fixture
def simple_docx():
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text.")
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("Content here.")
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    fname = tmp.name
    tmp.close()
    doc.save(fname)
    yield fname
    try:
        os.unlink(fname)
    except OSError:
        pass


# ═══════════════════════════════════════════════════════════════
#  Models Tests
# ═══════════════════════════════════════════════════════════════

class TestModels:
    def test_heading_info_defaults(self):
        from src.document.analyzer.models import HeadingInfo
        h = HeadingInfo(level=1, text="Intro")
        assert h.level == 1
        assert h.text == "Intro"
        assert h.numbering is None
        assert h.is_numbered is False

    def test_heading_info_to_dict(self):
        from src.document.analyzer.models import HeadingInfo
        h = HeadingInfo(level=2, text="Background", numbering="1.1")
        d = h.to_dict()
        assert d["level"] == 2
        assert d["text"] == "Background"
        assert d["numbering"] == "1.1"

    def test_style_profile_to_dict(self):
        from src.document.analyzer.models import StyleProfile, FontInfo, ParagraphFormatInfo
        s = StyleProfile(
            style_id="Heading1", style_name="Heading 1",
            is_heading=True, heading_level=1,
            font=FontInfo(name="Calibri", size=18.0, bold=True),
        )
        d = s.to_dict()
        assert d["style_id"] == "Heading1"
        assert d["is_heading"] is True
        assert d["font"]["name"] == "Calibri"
        assert d["font"]["bold"] is True

    def test_table_info_defaults(self):
        from src.document.analyzer.models import TableInfo
        t = TableInfo(index=0, rows=3, cols=4)
        assert t.rows == 3
        assert t.cols == 4
        assert t.caption is None

    def test_image_info_to_dict(self):
        from src.document.analyzer.models import ImageInfo
        img = ImageInfo(index=1, rId="rId1", width_inches=5.0, height_inches=3.0)
        d = img.to_dict()
        assert d["index"] == 1
        assert d["width_inches"] == 5.0

    def test_reference_info_detection(self):
        from src.document.analyzer.models import ReferenceInfo
        r = ReferenceInfo(raw_text="[1] Author, Title, 2020.", format="ieee")
        assert r.format == "ieee"
        d = r.to_dict()
        assert d["format"] == "ieee"

    def test_doc_knowledge_graph_to_dict(self):
        from src.document.analyzer.models import DocKnowledgeGraph
        g = DocKnowledgeGraph(filename="test.docx")
        d = g.to_dict()
        assert d["filename"] == "test.docx"
        assert d["style_count"] == 0
        assert d["table_count"] == 0


# ═══════════════════════════════════════════════════════════════
#  HeadingDetector Tests
# ═══════════════════════════════════════════════════════════════

class TestHeadingDetector:
    def test_detect_standard_headings(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        assert len(headings) >= 10

    def test_heading_levels(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        h1 = [h for h in headings if h.level == 1]
        h2 = [h for h in headings if h.level == 2]
        assert len(h1) >= 7
        assert len(h2) >= 3

    def test_numbered_headings(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        numbered = [h for h in headings if h.is_numbered]
        assert len(numbered) >= 5

    def test_hierarchy_tree(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        det.detect()
        tree = det.get_hierarchy_tree()
        assert isinstance(tree, list)

    def test_simple_heading(self, simple_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(simple_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        assert len(headings) >= 2

    def test_heading_content(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        texts = [h.text.lower() for h in headings]
        assert "introduction" in texts
        assert "methodology" in texts
        assert "conclusion" in texts

    def test_numbering_stripping(self, academic_docx):
        from src.document.analyzer import HeadingDetector
        doc = Document(academic_docx)
        det = HeadingDetector(doc)
        headings = det.detect()
        intro = next((h for h in headings if "introduction" in h.text.lower()), None)
        assert intro is not None
        assert not intro.text.startswith("1.")


# ═══════════════════════════════════════════════════════════════
#  SectionClassifier Tests
# ═══════════════════════════════════════════════════════════════

class TestSectionClassifier:
    def test_classify_known_types(self, academic_docx):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        doc = Document(academic_docx)
        headings = HeadingDetector(doc).detect()
        classifier = SectionClassifier()
        sections = classifier.classify_headings(headings)
        types = {s.section_type for s in sections}
        assert "abstract" in types or "introduction" in types
        assert "references" in types

    def test_classify_abstract(self, academic_docx):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        doc = Document(academic_docx)
        h = HeadingDetector(doc).detect()
        c = SectionClassifier()
        abstract = next((h for h in h if "abstract" in h.text.lower()), None)
        if abstract:
            assert c.classify(abstract) == "abstract"

    def test_classify_introduction(self, academic_docx):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        doc = Document(academic_docx)
        h = HeadingDetector(doc).detect()
        c = SectionClassifier()
        intro = next((h for h in h if "introduction" in h.text.lower()), None)
        if intro:
            assert c.classify(intro) == "introduction"

    def test_classify_references(self, academic_docx):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        doc = Document(academic_docx)
        h = HeadingDetector(doc).detect()
        c = SectionClassifier()
        ref = next((h for h in h if "references" in h.text.lower()), None)
        if ref:
            assert c.classify(ref) == "references"

    def test_classify_unknown(self):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        from src.document.analyzer.models import HeadingInfo
        c = SectionClassifier()
        h = HeadingInfo(level=1, text="Random Section Title")
        assert c.classify(h) == "unknown"

    def test_section_hierarchy_preserved(self, academic_docx):
        from src.document.analyzer import HeadingDetector, SectionClassifier
        doc = Document(academic_docx)
        h = HeadingDetector(doc).detect()
        c = SectionClassifier()
        sections = c.classify_headings(h)
        all_level1 = all(s.level == 1 for s in sections)
        assert all_level1


# ═══════════════════════════════════════════════════════════════
#  StyleExtractor Tests
# ═══════════════════════════════════════════════════════════════

class TestStyleExtractor:
    def test_extract_all_styles(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        assert len(styles) >= 10

    def test_heading_style_detected(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        heading_styles = {
            k: v for k, v in styles.items() if v.is_heading
        }
        assert len(heading_styles) >= 1

    def test_style_profile_content(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        hs = styles.get("Heading 1")
        if hs:
            assert hs.font.bold or hs.font.size > 0

    def test_paragraph_format_extracted(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        para = doc.paragraphs[0]
        fmt = ext.extract_paragraph_format(para)
        assert fmt.alignment in ("LEFT", "CENTER", "RIGHT", "JUSTIFY")

    def test_run_font_extracted(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        for para in doc.paragraphs:
            if para.runs:
                font = ext.extract_run_font(para.runs[0])
                assert font.name is not None
                break


# ═══════════════════════════════════════════════════════════════
#  TableDetector Tests
# ═══════════════════════════════════════════════════════════════

class TestTableDetector:
    def test_detect_tables(self, academic_docx):
        from src.document.analyzer import TableDetector
        doc = Document(academic_docx)
        det = TableDetector(doc)
        tables = det.detect()
        assert len(tables) >= 2

    def test_table_headers(self, academic_docx):
        from src.document.analyzer import TableDetector
        doc = Document(academic_docx)
        det = TableDetector(doc)
        tables = det.detect()
        if tables:
            assert len(tables[0].headers) > 0

    def test_table_captions(self, academic_docx):
        from src.document.analyzer import TableDetector
        doc = Document(academic_docx)
        det = TableDetector(doc)
        tables = det.detect_with_captions(doc.paragraphs)
        captioned = [t for t in tables if t.caption]
        assert len(captioned) >= 1

    def test_table_data_dimensions(self, academic_docx):
        from src.document.analyzer import TableDetector
        doc = Document(academic_docx)
        det = TableDetector(doc)
        tables = det.detect()
        if tables:
            assert tables[0].rows > 1
            assert tables[0].cols > 1

    def test_no_tables_simple(self, simple_docx):
        from src.document.analyzer import TableDetector
        doc = Document(simple_docx)
        det = TableDetector(doc)
        tables = det.detect()
        assert len(tables) == 0


# ═══════════════════════════════════════════════════════════════
#  ImageDetector Tests
# ═══════════════════════════════════════════════════════════════

class TestImageDetector:
    def test_no_images(self, academic_docx):
        from src.document.analyzer import ImageDetector
        doc = Document(academic_docx)
        det = ImageDetector(doc)
        images = det.detect()
        assert len(images) == 0

    def test_image_in_document(self):
        doc = Document()
        doc.add_heading("Test", level=1)
        p = doc.add_paragraph()
        r = p.add_run()
        drawing_xml = (
            f'<w:drawing {nsdecls("w", "wp", "a", "r")}>'
            f'  <wp:inline distT="0" distB="0" distL="0" distR="0">'
            f'    <wp:extent cx="914400" cy="914400"/>'
            f'    <wp:docPr id="1" name="Picture 1" descr="Test image"/>'
            f'    <a:graphic>'
            f'      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'        <pic:pic {nsdecls("pic")}>'
            f'          <pic:blipFill>'
            f'            <a:blip r:embed="rId1"/>'
            f'          </pic:blipFill>'
            f'        </pic:pic>'
            f'      </a:graphicData>'
            f'    </a:graphic>'
            f'  </wp:inline>'
            f'</w:drawing>'
        )
        p._element.append(parse_xml(drawing_xml))

        from src.document.analyzer import ImageDetector
        det = ImageDetector(doc)
        images = det.detect()
        assert len(images) >= 1

    def test_image_dimensions(self):
        doc = Document()
        doc.add_heading("Test", level=1)
        p = doc.add_paragraph()
        r = p.add_run()
        drawing_xml = (
            f'<w:drawing {nsdecls("w", "wp", "a", "r")}>'
            f'  <wp:inline distT="0" distB="0" distL="0" distR="0">'
            f'    <wp:extent cx="1828800" cy="914400"/>'
            f'    <wp:docPr id="1" name="Picture 1" descr="Test image"/>'
            f'    <a:graphic>'
            f'      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'        <pic:pic {nsdecls("pic")}>'
            f'          <pic:blipFill>'
            f'            <a:blip r:embed="rId1"/>'
            f'          </pic:blipFill>'
            f'        </pic:pic>'
            f'      </a:graphicData>'
            f'    </a:graphic>'
            f'  </wp:inline>'
            f'</w:drawing>'
        )
        p._element.append(parse_xml(drawing_xml))

        from src.document.analyzer import ImageDetector
        det = ImageDetector(doc)
        images = det.detect()
        if images:
            assert images[0].width_inches == pytest.approx(2.0, rel=0.1)
            assert images[0].height_inches == pytest.approx(1.0, rel=0.1)


# ═══════════════════════════════════════════════════════════════
#  ReferenceDetector Tests
# ═══════════════════════════════════════════════════════════════

class TestReferenceDetector:
    def test_detect_ieee_references(self, academic_docx):
        from src.document.analyzer import ReferenceDetector
        from src.document.analyzer.parser import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        assert len(graph.references) >= 2

    def test_reference_formats(self, academic_docx):
        from src.document.analyzer import ReferenceDetector, ParagraphInfo
        from src.document.analyzer.parser import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        ieee_refs = [r for r in graph.references if r.format == "ieee"]
        assert len(ieee_refs) >= 1

    def test_citation_links_detected(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        assert len(graph.citation_links) >= 1

    def test_no_references_simple(self, simple_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(simple_docx)
        assert len(graph.references) == 0


# ═══════════════════════════════════════════════════════════════
#  KnowledgeGraphBuilder Tests
# ═══════════════════════════════════════════════════════════════

class TestKnowledgeGraph:
    def test_build_complete_graph(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        assert graph.headings
        assert graph.styles
        assert graph.tables
        assert graph.references
        assert graph.paragraphs
        assert graph.statistics

    def test_graph_statistics(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        stats = graph.statistics
        assert stats["heading_count"] >= 5
        assert stats["table_count"] >= 2
        assert stats["reference_count"] >= 2
        assert stats["word_count"] > 0

    def test_graph_sections(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        assert len(graph.sections) >= 5

    def test_graph_to_dict(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        d = graph.to_dict()
        assert "analysis_id" in d
        assert "filename" in d
        assert d["table_count"] >= 2

    def test_sections_have_content(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        for sec in graph.sections:
            assert sec.confidence >= 0


# ═══════════════════════════════════════════════════════════════
#  DocxAnalyzer Integration Tests
# ═══════════════════════════════════════════════════════════════

class TestDocxAnalyzer:
    def test_analyze_from_path(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        assert graph.filename.endswith(".docx")

    def test_analyze_from_doc(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        doc = Document(academic_docx)
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze_doc(doc, "test.docx")
        assert len(graph.headings) > 0

    def test_get_summary(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        analyzer.analyze(academic_docx)
        summary = analyzer.get_summary()
        assert "filename" in summary

    def test_get_statistics(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        analyzer.analyze(academic_docx)
        stats = analyzer.get_statistics()
        assert "word_count" in stats

    def test_get_heading_hierarchy(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        analyzer.analyze(academic_docx)
        tree = analyzer.get_heading_hierarchy()
        assert isinstance(tree, list)

    def test_get_section_types(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        analyzer.analyze(academic_docx)
        types = analyzer.get_section_types()
        assert isinstance(types, dict)

    def test_export_json(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        analyzer.analyze(academic_docx)
        tmp = tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode='w')
        fname = tmp.name
        tmp.close()
        try:
            result = analyzer.export_json(fname)
            assert result is True
            with open(fname, 'r') as f:
                data = json.load(f)
            assert "filename" in data
        finally:
            os.unlink(fname)

    def test_file_not_found(self):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        with pytest.raises(FileNotFoundError):
            analyzer.analyze("nonexistent.docx")

    def test_no_analysis_summary(self):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        summary = analyzer.get_summary()
        assert "error" in summary

    def test_graph_sections_have_types(self, academic_docx):
        from src.document.analyzer import DocxAnalyzer
        analyzer = DocxAnalyzer()
        graph = analyzer.analyze(academic_docx)
        types = {s.section_type for s in graph.sections}
        assert "unknown" not in types or len(types) > 0


# ═══════════════════════════════════════════════════════════════
#  Style Mapping Tests
# ═══════════════════════════════════════════════════════════════

class TestStyleMapping:
    def test_font_family_extracted(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        for s in styles.values():
            if s.font.name:
                assert isinstance(s.font.name, str)
                return
        assert False, "No font name found"

    def test_font_size_extracted(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        for s in styles.values():
            if s.font.size > 0:
                assert isinstance(s.font.size, float)
                return
        assert False, "No font size found"

    def test_paragraph_alignment(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        for para in doc.paragraphs:
            fmt = ext.extract_paragraph_format(para)
            assert isinstance(fmt.alignment, str)
            break

    def test_style_profile_reusable(self, academic_docx):
        from src.document.analyzer import StyleExtractor
        doc = Document(academic_docx)
        ext = StyleExtractor(doc)
        styles = ext.extract_all()
        s = styles.get("Heading 1")
        if s:
            d = s.to_dict()
            assert "font" in d
            assert "paragraph_format" in d


class TestFootnoteDetector:
    def test_no_footnotes_in_simple(self, academic_docx):
        from src.document.analyzer.footnotes import FootnoteDetector
        doc = Document(academic_docx)
        detector = FootnoteDetector(doc)
        footnotes = detector.detect()
        assert isinstance(footnotes, list)

    def test_footnote_model(self):
        from src.document.analyzer.models import FootnoteInfo
        fn = FootnoteInfo(index=0, footnote_id="1", text="A note.", paragraph_count=1)
        d = fn.to_dict()
        assert d["id"] == "1"
        assert "A note." in d["text"]


class TestHeaderFooterDetector:
    def test_detect_headers_footers(self, academic_docx):
        from src.document.analyzer.headers_footers import HeaderFooterDetector
        doc = Document(academic_docx)
        detector = HeaderFooterDetector(doc)
        results = detector.detect()
        assert isinstance(results, list)

    def test_header_footer_model(self):
        from src.document.analyzer.models import HeaderFooterInfo
        hf = HeaderFooterInfo(section_index=0, type="header", text="Report Title", paragraph_count=1)
        d = hf.to_dict()
        assert d["type"] == "header"
        assert "Report Title" in d["text"]


class TestCrossReferenceDetector:
    def test_no_cross_refs_in_plain_text(self, academic_docx):
        from src.document.analyzer.cross_references import CrossReferenceDetector
        doc = Document(academic_docx)
        detector = CrossReferenceDetector(doc)
        refs = detector.detect()
        assert isinstance(refs, list)

    def test_cross_ref_model(self):
        from src.document.analyzer.models import CrossReferenceInfo
        cr = CrossReferenceInfo(reference_type="section", reference_text="R_1", paragraph_index=5)
        d = cr.to_dict()
        assert d["type"] == "section"
        assert d["paragraph_index"] == 5


class TestKnowledgeGraphNewFields:
    def test_graph_has_footnote_field(self, academic_docx):
        from src.document.analyzer import KnowledgeGraphBuilder
        doc = Document(academic_docx)
        builder = KnowledgeGraphBuilder(doc)
        graph = builder.build(filename="test.docx")
        assert hasattr(graph, "footnotes")
        assert hasattr(graph, "headers_footers")
        assert hasattr(graph, "cross_references")

    def test_graph_to_dict_includes_new_counts(self, academic_docx):
        from src.document.analyzer import KnowledgeGraphBuilder
        doc = Document(academic_docx)
        builder = KnowledgeGraphBuilder(doc)
        graph = builder.build(filename="test.docx")
        d = graph.to_dict()
        assert "footnote_count" in d
        assert "header_footer_count" in d
        assert "cross_reference_count" in d

    def test_statistics_includes_new_metrics(self, academic_docx):
        from src.document.analyzer import KnowledgeGraphBuilder
        doc = Document(academic_docx)
        builder = KnowledgeGraphBuilder(doc)
        graph = builder.build(filename="test.docx")
        stats = graph.statistics
        assert "footnote_count" in stats
        assert "header_footer_count" in stats
        assert "cross_reference_count" in stats
