import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class MockProvider:
    """Mock LLM provider for deterministic testing without Ollama."""

    def __init__(self, model="mock-model", temperature=0.7, top_p=0.9, timeout=120):
        self.model = model
        self.temperature = temperature
        self.top_p = top_p
        self.timeout = timeout
        self._available = True
        self.chat_history = []

    def is_available(self):
        return self._available

    def set_available(self, available: bool):
        self._available = available

    def chat(self, messages, options=None):
        from src.providers.base import LLMResponse
        self.chat_history.append(messages)
        return LLMResponse(
            content='{"title": "Mock Report", "sections": [{"heading": "Mock", "content": "Mock content."}]}',
            model=self.model,
        )

    def generate(self, prompt, options=None):
        from src.providers.base import LLMResponse
        return LLMResponse(
            content="Mock generated content for testing purposes.",
            model=self.model,
        )

    def prepare_messages(self, system, user):
        from src.providers.base import Message
        return [Message(role="system", content=system), Message(role="user", content=user)]

    def get_provider_type(self):
        from src.providers.base import ProviderType
        return ProviderType.OLLAMA

    def __repr__(self):
        return f"MockProvider(model={self.model})"


@pytest.fixture
def mock_provider():
    """Fixture providing a MockProvider instance."""
    return MockProvider()


@pytest.fixture
def mock_provider_unavailable():
    """Fixture providing an unavailable MockProvider."""
    p = MockProvider()
    p.set_available(False)
    return p


@pytest.fixture
def sample_docx():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Test Report")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("A sample test document")
    subtitle_run.font.size = Pt(16)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("1. Executive Summary")
    doc.add_paragraph("2. Introduction")
    doc.add_paragraph("3. Analysis")
    doc.add_paragraph("4. Conclusion")

    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This is the executive summary of the test report. "
        "It provides an overview of the key findings and recommendations."
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is the introduction section. "
        "It provides background information on the topic."
    )
    doc.add_paragraph("This is a second paragraph in the introduction.")

    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "This subsection describes the methodology used in the report."
    )

    doc.add_heading("Analysis", level=1)
    doc.add_paragraph(
        "This section contains the detailed analysis."
    )

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for i, (h, d) in enumerate([("Metric", "Value"), ("Revenue", "$100K"), ("Cost", "$60K")]):
        table.rows[i].cells[0].text = h
        table.rows[i].cells[1].text = d

    doc.add_paragraph(
        "Additional analysis content after the table."
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "This is the conclusion section summarizing the report findings."
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    yield temp_path

    try:
        os.unlink(temp_path)
    except OSError:
        pass


@pytest.fixture
def sample_document():
    from docx import Document
    return Document()


@pytest.fixture
def build_tree():
    from src.document.structure.model import build_tree
    return build_tree


@pytest.fixture
def SectionLocator():
    from src.document.structure.locator import SectionLocator
    return SectionLocator


@pytest.fixture
def EditingPlanner():
    from src.document.structure.planner import EditingPlanner
    return EditingPlanner
