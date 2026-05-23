"""
Document Generator Module
=========================
Generates Word documents using python-docx.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_DOCX = "output.docx"


def set_font(run, font_name, size, color=None, bold=False):
    """Set font properties for a text run."""
    run.font.name = font_name
    run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold


def generate_document(content: dict = None, output_path: str = OUTPUT_DOCX) -> bool:
    """
    Generate Word document with content.
    
    Args:
        content: Dictionary with report content
        output_path: Path to save the document
        
    Returns:
        bool: Success status
    """
    if content is None:
        content = {}
    
    # Use defaults if any fields are missing
    title = content.get('title', 'Report')
    subtitle = content.get('subtitle', 'Technical Report')
    author = content.get('author', 'AI Generator')
    date = content.get('date', 'May 2026')
    toc_entries = content.get('toc_entries', ['Section 1', 'Section 2'])
    exec_summary = content.get('executive_summary', 'No summary available.')
    intro = content.get('introduction', 'No introduction available.')
    threats = content.get('threats', [])
    ml_methods = content.get('ml_methods', [])
    case_studies = content.get('case_studies', 'No case studies.')
    future_trends = content.get('future_trends', [])
    conclusion = content.get('conclusion', 'No conclusion.')
    
    print("\n" + "=" * 50)
    print("[GENERATING] Creating Word document...")
    print("=" * 50)
    
    print(f"[INFO] Generating report: {title}")
    
    try:
        doc = Document()
        
        # Set page margins
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Cover page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        set_font(title_run, 'Times New Roman', Pt(28), RGBColor(0, 51, 51), True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        set_font(subtitle_run, 'Calibri', Pt(16), RGBColor(100, 100, 100))
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f"Author: {author}\nDate: {date}")
        set_font(info_run, 'Calibri', Pt(12), RGBColor(80, 80, 80))
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of contents
        doc.add_heading('Table of Contents', level=1)
        for i, entry in enumerate(toc_entries, 1):
            p = doc.add_paragraph(f'{i}. {entry}')
            p.paragraph_format.space_before = Pt(6)
        
        doc.add_page_break()
        
        # Content sections
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(exec_summary)
        
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(intro)
        
        # Main content section
        doc.add_heading('Analysis', level=1)
        
        # Threats/Items table
        if threats:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Topic'
            hdr_cells[1].text = 'Description'
            for item, desc in threats:
                row = table.add_row().cells
                row[0].text = str(item)
                row[1].text = str(desc)
        
        # Methods
        if ml_methods:
            doc.add_heading('Key Methods', level=2)
            for method in ml_methods:
                p = doc.add_paragraph(str(method))
                p.style = 'List Bullet'
        
        # Case Studies
        if case_studies:
            doc.add_heading('Case Studies', level=1)
            doc.add_paragraph(str(case_studies))
        
        # Future Trends
        if future_trends:
            doc.add_heading('Future Trends', level=1)
            for trend in future_trends:
                p = doc.add_paragraph(str(trend))
                p.style = 'List Bullet'
        
        # Conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(conclusion)
        
        # Save
        doc.save(output_path)
        
        if os.path.exists(output_path):
            size = os.path.getsize(output_path)
            print(f"[OK] Document created: {output_path}")
            print(f"     Size: {size / 1024:.1f} KB")
            return True
        return False
            
    except Exception as e:
        print(f"[ERROR] Failed to generate document: {e}")
        import traceback
        traceback.print_exc()
        return False


def verify_document(output_path: str = OUTPUT_DOCX) -> dict:
    """Verify the generated document."""
    if not os.path.exists(output_path):
        print(f"[ERROR] {output_path} not found!")
        return None
    
    size = os.path.getsize(output_path)
    print(f"[OK] File exists: {output_path}")
    print(f"     Size: {size / 1024:.1f} KB")
    
    try:
        doc = Document(output_path)
        print(f"     Paragraphs: {len(doc.paragraphs)}")
        print(f"     Tables: {len(doc.tables)}")
    except Exception as e:
        print(f"     [INFO] Could not read details: {e}")
    
    return {'path': output_path, 'size': size}


if __name__ == "__main__":
    generate_document()
    verify_document()