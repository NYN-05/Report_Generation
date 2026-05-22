"""
Document Parser Module
======================
Parses and analyzes DOCX documents.
"""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.table import Table

from .base import DocumentMetadata
from src.core.logger import get_logger
from src.core.exceptions import DocumentParseError

logger = get_logger(__name__)


class DocumentParser:
    """Parses DOCX documents and extracts information."""

    def __init__(self, document_path: str):
        self.document_path = document_path
        self._document: Optional[Document] = None

    def parse(self) -> Dict[str, Any]:
        """Parse the document and return structured data."""
        if not os.path.exists(self.document_path):
            raise DocumentParseError(f"Document not found: {self.document_path}")

        try:
            self._document = Document(self.document_path)
            return self._extract_content()
        except Exception as e:
            raise DocumentParseError(f"Failed to parse document: {e}")

    def _extract_content(self) -> Dict[str, Any]:
        """Extract content from document."""
        return {
            "metadata": self._extract_metadata(),
            "paragraphs": self._extract_paragraphs(),
            "headings": self._extract_headings(),
            "tables": self._extract_tables(),
            "styles": self._extract_styles(),
        }

    def _extract_metadata(self) -> DocumentMetadata:
        """Extract document metadata."""
        core_props = self._document.core_properties

        return DocumentMetadata(
            title=core_props.title or "",
            author=core_props.author or "",
            date=str(core_props.modified or ""),
            paragraph_count=len(self._document.paragraphs),
            table_count=len(self._document.tables)
        )

    def _extract_paragraphs(self) -> List[Dict[str, str]]:
        """Extract all paragraphs."""
        paragraphs = []

        for para in self._document.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })

        return paragraphs

    def _extract_headings(self) -> List[Dict[str, Any]]:
        """Extract all headings."""
        headings = []

        for para in self._document.paragraphs:
            if para.style.name.startswith('Heading'):
                level = 1
                if len(para.style.name) > 7:
                    try:
                        level = int(para.style.name[-1])
                    except:
                        pass

                headings.append({
                    "text": para.text,
                    "level": level,
                    "style": para.style.name
                })

        return headings

    def _extract_tables(self) -> List[Dict[str, Any]]:
        """Extract all tables."""
        tables = []

        for idx, table in enumerate(self._document.tables):
            table_data = []

            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            tables.append({
                "index": idx,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": table_data
            })

        return tables

    def _extract_styles(self) -> List[str]:
        """Extract all used styles."""
        styles = set()

        for para in self._document.paragraphs:
            if para.style:
                styles.add(para.style.name)

        return sorted(list(styles))

    def get_text_content(self) -> str:
        """Get all text content as a single string."""
        if not self._document:
            self._document = Document(self.document_path)

        return "\n".join(para.text for para in self._document.paragraphs if para.text.strip())

    def get_section_text(self, heading: str) -> Optional[str]:
        """Get text content under a specific heading."""
        if not self._document:
            self._document = Document(self.document_path)

        capture = False
        section_text = []

        for para in self._document.paragraphs:
            if para.style.name.startswith('Heading'):
                if heading.lower() in para.text.lower():
                    capture = True
                else:
                    capture = False

            if capture and para.text.strip():
                section_text.append(para.text)

        return "\n".join(section_text) if section_text else None

    def count_words(self) -> int:
        """Count total words in document."""
        text = self.get_text_content()
        return len(text.split())


def parse_document(path: str) -> Dict[str, Any]:
    """Convenience function to parse a document."""
    parser = DocumentParser(path)
    return parser.parse()