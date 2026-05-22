"""
Placeholder Handler Module
==========================
Handles placeholder replacement in templates.
"""

import re
from typing import Dict, Any, List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from src.core.logger import get_logger
from src.core.exceptions import PlaceholderError

logger = get_logger(__name__)


class PlaceholderHandler:
    """Handles placeholder detection and replacement."""

    PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')

    def __init__(self, document: Document):
        self.document = document
        self._replacements_made = 0

    def replace_all(self, mapping: Dict[str, str]) -> int:
        """Replace all placeholders with content from mapping."""
        self._replacements_made = 0

        for para in self.document.paragraphs:
            self._replace_in_paragraph(para, mapping)

        logger.info(f"Replaced {self._replacements_made} placeholders")
        return self._replacements_made

    def _replace_in_paragraph(self, paragraph: Paragraph, mapping: Dict[str, str]):
        """Replace placeholders in a single paragraph."""
        text = paragraph.text

        if not self.PLACEHOLDER_PATTERN.search(text):
            return

        for run in paragraph.runs:
            run_text = run.text
            new_text = run_text

            for placeholder, content in mapping.items():
                pattern = re.compile(r'\{\{' + placeholder + r'\}\}')
                new_text = pattern.sub(content, new_text)

            if new_text != run_text:
                run.text = new_text
                self._replacements_made += 1

    def replace_single(self, placeholder: str, content: str) -> bool:
        """Replace a single placeholder."""
        found = False

        for para in self.document.paragraphs:
            for run in para.runs:
                if f"{{{{{placeholder}}}}}" in run.text:
                    run.text = run.text.replace(f"{{{{{placeholder}}}}}", content)
                    self._replacements_made += 1
                    found = True

        return found

    def find_placeholders(self) -> List[str]:
        """Find all unique placeholder names."""
        placeholders = set()

        for para in self.document.paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(para.text)
            placeholders.update(matches)

        return sorted(list(placeholders))

    def validate_mapping(self, mapping: Dict[str, str]) -> Dict[str, bool]:
        """Validate that all required placeholders have values."""
        required = self.find_placeholders()
        validation = {}

        for placeholder in required:
            validation[placeholder] = placeholder in mapping

        return validation


class PlaceholderMatcher:
    """Matches content to template placeholders."""

    def __init__(self):
        self._common_mappings = {
            "title": "title",
            "subtitle": "subtitle",
            "author": "author",
            "date": "date",
            "summary": "executive_summary",
            "summary_text": "executive_summary",
            "intro": "introduction",
            "introduction_text": "introduction",
            "conclusion_text": "conclusion",
            "concl": "conclusion",
        }

    def create_mapping(self, content: Dict[str, Any], placeholders: List[str]) -> Dict[str, str]:
        """Create mapping from content to placeholders."""
        mapping = {}

        for placeholder in placeholders:
            placeholder_lower = placeholder.lower()

            if placeholder_lower in content:
                mapping[placeholder] = str(content[placeholder_lower])

            for key, content_key in self._common_mappings.items():
                if placeholder_lower == key and content_key in content:
                    mapping[placeholder] = str(content[content_key])

        return mapping

    def suggest_fill(self, placeholders: List[str], existing_content: Dict[str, Any]) -> Dict[str, str]:
        """Suggest content for placeholders based on available content."""
        suggestions = {}

        for placeholder in placeholders:
            suggested = self._find_suggestion(placeholder, existing_content)
            if suggested:
                suggestions[placeholder] = suggested

        return suggestions

    def _find_suggestion(self, placeholder: str, content: Dict[str, Any]) -> Optional[str]:
        """Find a suggestion for a placeholder."""
        placeholder_lower = placeholder.lower()

        for key, value in content.items():
            if placeholder_lower in key.lower():
                return str(value)

        return None


def replace_placeholders(document: Document, mapping: Dict[str, str]) -> int:
    """Convenience function to replace placeholders."""
    handler = PlaceholderHandler(document)
    return handler.replace_all(mapping)