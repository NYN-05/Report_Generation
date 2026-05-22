"""
Template Analyzer Module
========================
Analyzes template structure and placeholders.
"""

import re
from dataclasses import dataclass
from typing import Dict, List, Any, Optional
from docx import Document

from src.core.logger import get_logger

logger = get_logger(__name__)


@dataclass
class PlaceholderInfo:
    """Information about a placeholder."""
    name: str
    pattern: str
    context: str
    suggested_type: str


class TemplateAnalyzer:
    """Analyzes DOCX template structure."""

    def __init__(self, document: Document):
        self.document = document
        self._placeholders: List[PlaceholderInfo] = []
        self._heading_positions: Dict[str, int] = {}

    def analyze(self) -> Dict[str, Any]:
        """Analyze the template and extract structure."""
        return {
            "placeholders": self.extract_placeholders(),
            "headings": self._extract_headings(),
            "sections": self._extract_sections(),
            "tables": self._analyze_tables(),
            "styles": self._extract_styles()
        }

    def extract_placeholders(self) -> List[Dict[str, str]]:
        """Extract placeholders from document."""
        placeholders = []
        pattern = re.compile(r'\{\{(\w+)\}\}')

        for para in self.document.paragraphs:
            matches = pattern.findall(para.text)
            for match in matches:
                placeholders.append({
                    "name": match,
                    "full_match": f"{{{{{match}}}}}",
                    "context": para.text[:100]
                })

        return placeholders

    def _extract_headings(self) -> List[Dict[str, Any]]:
        """Extract heading structure."""
        headings = []

        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                level = 1
                try:
                    level = int(para.style.name[-1])
                except:
                    pass

                headings.append({
                    "text": para.text,
                    "level": level,
                    "position": len(headings)
                })

        return headings

    def _extract_sections(self) -> List[Dict[str, Any]]:
        """Extract document sections."""
        sections = []
        current_section = None

        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                if current_section:
                    sections.append(current_section)

                current_section = {
                    "title": para.text,
                    "level": int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1,
                    "content": []
                }
            elif current_section and para.text.strip():
                current_section["content"].append(para.text)

        if current_section:
            sections.append(current_section)

        return sections

    def _analyze_tables(self) -> List[Dict[str, Any]]:
        """Analyze tables in template."""
        tables = []

        for idx, table in enumerate(self.document.tables):
            table_info = {
                "index": idx,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "has_headers": len(table.rows) > 0
            }
            tables.append(table_info)

        return tables

    def _extract_styles(self) -> List[str]:
        """Extract all styles used in template."""
        styles = set()

        for para in self.document.paragraphs:
            if para.style:
                styles.add(para.style.name)

        return sorted(list(styles))

    def get_schema(self) -> Dict[str, Any]:
        """Get template schema for content generation."""
        placeholders = self.extract_placeholders()
        headings = self._extract_headings()

        schema = {
            "required_fields": [p["name"] for p in placeholders],
            "section_order": [h["text"] for h in headings if h["level"] == 1],
            "structure": self._extract_sections()
        }

        return schema

    def suggest_placeholder_mapping(self, content: Dict) -> Dict[str, str]:
        """Suggest mapping between content and placeholders."""
        placeholders = {p["name"].lower(): p["name"] for p in self.extract_placeholders()}
        mapping = {}

        content_keys = {
            "title": "title",
            "subtitle": "subtitle",
            "author": "author",
            "date": "date",
            "executive_summary": "executive_summary",
            "introduction": "introduction",
            "conclusion": "conclusion"
        }

        for content_key, placeholder in content_keys.items():
            if content_key in placeholders:
                mapping[content_key] = placeholders[content_key]

        return mapping