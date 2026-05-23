import os
from typing import List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .models import ReportPlan, PlanSection
from src.document.builder import DocumentBuilder
from src.core.logger import get_logger
from src.document.styles import StyleManager

logger = get_logger(__name__)


class BlueprintBuilder:
    """Generates DOCX documents from a ReportPlan."""

    def __init__(self):
        self._doc: Optional[Document] = None
        self._styles = StyleManager.get_instance()

    def build(self, plan: ReportPlan, output_path: str = "output.docx") -> bool:
        builder = DocumentBuilder()
        builder.create()
        self._builder = builder
        self._doc = builder._document

        # 1. Cover page
        if any(s.blueprint_section_id == "cover_page" for s in plan.sections):
            self._add_cover_page(plan)
        else:
            self._add_title_page(plan)

        # 2. Preliminary sections
        for section in plan.sections:
            bp_id = section.blueprint_section_id
            if bp_id in ("cover_page", "table_of_contents", "list_of_figures", "list_of_tables"):
                continue
            if bp_id in ("certificate", "declaration", "acknowledgement", "abstract"):
                self._add_preliminary_section(section)

        # 3. Table of Contents
        has_toc = any(s.blueprint_section_id == "table_of_contents" for s in plan.sections)
        if has_toc:
            self._add_toc()

        # 4. List of Figures
        has_lof = any(s.blueprint_section_id == "list_of_figures" for s in plan.sections)
        if has_lof and plan.total_figures > 0:
            self._add_list_of_figures(plan)

        # 5. List of Tables
        has_lot = any(s.blueprint_section_id == "list_of_tables" for s in plan.sections)
        if has_lot and plan.total_tables > 0:
            self._add_list_of_tables(plan)

        # 6. Main content sections
        for section in plan.sections:
            bp_id = section.blueprint_section_id
            if bp_id in ("cover_page", "certificate", "declaration", "acknowledgement",
                         "abstract", "table_of_contents", "list_of_figures", "list_of_tables"):
                continue
            if bp_id == "chapters" or (section.level >= 1 and section.content):
                self._add_content_section(section)

        # 7. References
        ref_section = self._find_section(plan.sections, "references")
        if ref_section:
            self._add_references(ref_section, plan)

        # 8. Appendices
        app_section = self._find_section(plan.sections, "appendices")
        if app_section:
            self._add_appendix(app_section)

        # Save
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self._doc.save(output_path)
        logger.info(f"Blueprint document saved: {output_path}")
        return True

    def _add_cover_page(self, plan: ReportPlan):
        self._styles.setup_document(self._doc)

        for _ in range(6):
            self._doc.add_paragraph()

        s = self._styles.get_styles()
        title_p = self._doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._styles.write_run(title_p, plan.title, s.cover_page.title_font)

        if plan.subtitle:
            sub_p = self._doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(sub_p, plan.subtitle, s.cover_page.subtitle_font)

        self._doc.add_paragraph()

        info_lines = []
        if plan.author:
            info_lines.append(f"Author: {plan.author}")
        if plan.date:
            info_lines.append(f"Date: {plan.date}")
        if info_lines:
            info_p = self._doc.add_paragraph()
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(info_p, "\n".join(info_lines), s.cover_page.author_font)

        self._doc.add_page_break()

    def _add_title_page(self, plan: ReportPlan):
        self._doc.add_heading(plan.title, level=0)
        if plan.subtitle:
            self._doc.add_paragraph(plan.subtitle)
        if plan.author:
            self._doc.add_paragraph(f"Author: {plan.author}")
        if plan.date:
            self._doc.add_paragraph(f"Date: {plan.date}")
        self._doc.add_page_break()

    def _add_preliminary_section(self, section: PlanSection):
        self._doc.add_heading(section.heading, level=1)
        if section.content:
            for para_text in section.content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    self._doc.add_paragraph(para_text)

    def _add_toc(self):
        self._doc.add_page_break()
        toc_heading = self._doc.add_heading("Table of Contents", level=1)

        p = self._doc.add_paragraph()
        run = p.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar_begin)

        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar_end)

    def _add_list_of_figures(self, plan: ReportPlan):
        self._doc.add_page_break()
        self._doc.add_heading("List of Figures", level=1)
        figure_descriptions = []
        for section in plan.sections:
            if section.figure_description:
                figure_descriptions.append(section.figure_description)
        for i in range(1, plan.total_figures + 1):
            desc = figure_descriptions[i - 1] if i <= len(figure_descriptions) else f"Figure illustrating {plan.title.lower()}"
            p = self._doc.add_paragraph(f"Figure {i}: {desc}")

    def _add_list_of_tables(self, plan: ReportPlan):
        self._doc.add_page_break()
        self._doc.add_heading("List of Tables", level=1)
        table_descriptions = []
        for section in plan.sections:
            if section.table_headers:
                table_descriptions.append(f"Table showing {', '.join(section.table_headers[:3])}")
        for i in range(1, plan.total_tables + 1):
            desc = table_descriptions[i - 1] if i <= len(table_descriptions) else f"Table of {plan.title.lower()} data"
            p = self._doc.add_paragraph(f"Table {i}: {desc}")

    def _add_content_section(self, section: PlanSection):
        if section.content:
            self._doc.add_heading(section.heading, level=section.level)
            for para_text in section.content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    self._doc.add_paragraph(para_text)

        for subsection in section.subsections:
            if subsection.content:
                self._doc.add_heading(subsection.heading, level=subsection.level)
                for para_text in subsection.content.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        self._doc.add_paragraph(para_text)

            if subsection.requires_table and subsection.table_headers:
                self._add_table(subsection)

    def _add_table(self, section: PlanSection):
        headers = section.table_headers
        rows = section.table_rows
        table = self._doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                if ci < len(headers):
                    table.rows[ri + 1].cells[ci].text = cell

        self._doc.add_paragraph()

    def _add_references(self, section: PlanSection, plan: ReportPlan):
        self._doc.add_page_break()
        self._doc.add_heading("References", level=1)
        s = self._styles.get_styles()
        for ref in plan.references:
            p = self._doc.add_paragraph(ref)
            self._styles.apply_paragraph_style(p, s.reference)
            p.paragraph_format.first_line_indent = Inches(-s.reference.hanging_indent)
            p.paragraph_format.left_indent = Inches(s.reference.hanging_indent)

    def _add_appendix(self, section: PlanSection):
        self._doc.add_page_break()
        self._doc.add_heading("Appendices", level=1)
        if section.content:
            for para_text in section.content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    self._doc.add_paragraph(para_text)
        else:
            self._doc.add_paragraph("Appendix A: Supplementary Materials")
            self._doc.add_paragraph("Detailed technical specifications, code listings, or additional data tables relevant to the report.")

    def _find_section(self, sections: List[PlanSection], bp_id: str) -> Optional[PlanSection]:
        for s in sections:
            if s.blueprint_section_id == bp_id:
                return s
        return None
