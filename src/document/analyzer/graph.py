from typing import Dict, Any, List, Optional
from docx import Document as DocxDocument

from .models import (
    DocKnowledgeGraph, SectionInfo, HeadingInfo, StyleProfile,
    TableInfo, ImageInfo, ReferenceInfo, CitationLink, ParagraphInfo,
    FootnoteInfo, HeaderFooterInfo, CrossReferenceInfo,
    WatermarkInfo, EquationInfo,
)
from .heading import HeadingDetector
from .classifier import SectionClassifier
from .styles import StyleExtractor
from .tables import TableDetector
from .images import ImageDetector
from .references import ReferenceDetector
from .footnotes import FootnoteDetector
from .headers_footers import HeaderFooterDetector
from .cross_references import CrossReferenceDetector
from .watermarks import WatermarkDetector
from .equations import EquationDetector


class KnowledgeGraphBuilder:
    """Builds a complete document knowledge graph from analyzed components."""

    def __init__(self, doc: DocxDocument):
        self.doc = doc
        self.heading_detector = HeadingDetector(doc)
        self.classifier = SectionClassifier()
        self.style_extractor = StyleExtractor(doc)
        self.table_detector = TableDetector(doc)
        self.image_detector = ImageDetector(doc)
        self.reference_detector = ReferenceDetector()
        self.footnote_detector = FootnoteDetector(doc)
        self.header_footer_detector = HeaderFooterDetector(doc)
        self.cross_reference_detector = CrossReferenceDetector(doc)
        self.watermark_detector = WatermarkDetector(doc)
        self.equation_detector = EquationDetector(doc)

    def build(self, filename: str = "") -> DocKnowledgeGraph:
        headings = self.heading_detector.detect()
        styles = self.style_extractor.extract_all()
        tables = self.table_detector.detect_with_captions(self.doc.paragraphs)
        images = self.image_detector.detect(headings)
        sections = self.classifier.classify_headings(headings)

        paragraphs = self._extract_all_paragraphs()

        refs, citations = self.reference_detector.detect(
            paragraphs, self.doc.paragraphs
        )

        footnotes = self.footnote_detector.detect()
        headers_footers = self.header_footer_detector.detect()
        cross_references = self.cross_reference_detector.detect()
        watermarks = self.watermark_detector.detect()
        equations = self.equation_detector.detect()

        self._assign_content_to_sections(sections, tables, images, refs, headings)
        self._assign_tables_to_sections(sections, tables)
        self._assign_images_to_sections(sections, images)

        stats = self._compute_statistics(
            sections, headings, styles, tables, images, refs, citations,
            paragraphs, footnotes, headers_footers, cross_references,
            watermarks, equations,
        )

        graph = DocKnowledgeGraph(
            filename=filename or getattr(self.doc, 'name', 'document.docx'),
            sections=sections,
            headings=headings,
            styles=styles,
            tables=tables,
            figures=images,
            references=refs,
            citation_links=citations,
            paragraphs=paragraphs,
            footnotes=footnotes,
            headers_footers=headers_footers,
            cross_references=cross_references,
            watermarks=watermarks,
            equations=equations,
            statistics=stats,
        )
        return graph

    def _extract_all_paragraphs(self) -> List[ParagraphInfo]:
        results: List[ParagraphInfo] = []
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            pinfo = ParagraphInfo(
                text=text,
                style_name=style_name,
                paragraph_index=idx,
                font=self.style_extractor.extract_run_font(
                    para.runs[0] if para.runs else type('_run', (), {})()
                ),
                paragraph_format=self.style_extractor.extract_paragraph_format(para),
            )
            results.append(pinfo)
        return results

    def _assign_content_to_sections(
        self,
        sections: List[SectionInfo],
        tables: List[TableInfo],
        images: List[ImageInfo],
        refs: List[ReferenceInfo],
        headings: List[HeadingInfo],
    ):
        if not sections:
            return
        heading_para_indices = {h.paragraph_index: h for h in headings}
        heading_boundaries = sorted(heading_para_indices.keys())

        section_by_start = {}
        for sec in sections:
            if sec.heading:
                section_by_start[sec.heading.paragraph_index] = sec

        for pi in range(len(self.doc.paragraphs)):
            section = self._find_section_for_paragraph(
                pi, heading_boundaries, section_by_start, sections
            )
            if section is not None:
                para = self.doc.paragraphs[pi]
                text = para.text.strip()
                if text:
                    pinfo = ParagraphInfo(
                        text=text,
                        style_name=para.style.name if para.style else "Normal",
                        paragraph_index=pi,
                    )
                    section.paragraphs.append(pinfo)

    def _find_section_for_paragraph(
        self,
        pi: int,
        boundaries: List[int],
        section_by_start: Dict[int, SectionInfo],
        sections: List[SectionInfo],
    ) -> Optional[SectionInfo]:
        if not boundaries:
            return sections[0] if sections else None
        if pi < boundaries[0]:
            return sections[0] if sections else None
        last_start = boundaries[0]
        for b in boundaries:
            if pi < b:
                break
            last_start = b
        return section_by_start.get(last_start)

    def _assign_tables_to_sections(self, sections: List[SectionInfo], tables: List[TableInfo]):
        for sec in sections:
            if sec.heading:
                hw = sec.heading.text.lower()
                for t in tables:
                    if t.caption and hw in t.caption.lower():
                        sec.tables.append(t)
                self._assign_by_nearby_text(sec, tables, is_table=True)

    def _assign_images_to_sections(self, sections: List[SectionInfo], images: List[ImageInfo]):
        for img in images:
            if img.anchor_section:
                for sec in sections:
                    if sec.heading and sec.heading.text == img.anchor_section:
                        sec.images.append(img)
                        break

    def _assign_by_nearby_text(self, section: SectionInfo, items, is_table: bool = True):
        if not section.heading:
            return
        hpi = section.heading.paragraph_index
        assigned = {getattr(t, 'index', id(t)) for t in
                    (section.tables if is_table else section.images)}
        for item in items:
            item_id = getattr(item, 'index', id(item))
            if item_id in assigned:
                continue
            if is_table:
                seen_texts = {p.text.strip().lower() for p in section.paragraphs}
                if item.data and any(
                    any(cell.lower() in seen_texts for cell in row)
                    for row in item.data[:3]
                ):
                    section.tables.append(item)

    def _compute_statistics(
        self,
        sections: List[SectionInfo],
        headings: List[HeadingInfo],
        styles: Dict[str, StyleProfile],
        tables: List[TableInfo],
        images: List[ImageInfo],
        refs: List[ReferenceInfo],
        citations: List[CitationLink],
        paragraphs: List[ParagraphInfo],
        footnotes: List[FootnoteInfo] = None,
        headers_footers: List[HeaderFooterInfo] = None,
        cross_references: List[CrossReferenceInfo] = None,
        watermarks: List[WatermarkInfo] = None,
        equations: List[EquationInfo] = None,
    ) -> Dict[str, Any]:
        total_paras = len(self.doc.paragraphs)
        text_paras = sum(1 for p in self.doc.paragraphs if p.text.strip())
        word_count = sum(len(p.text.split()) for p in self.doc.paragraphs if p.text)
        char_count = sum(len(p.text) for p in self.doc.paragraphs if p.text)

        h1 = sum(1 for h in headings if h.level == 1)
        h2 = sum(1 for h in headings if h.level == 2)
        h3 = sum(1 for h in headings if h.level >= 3)

        level_counts = {}
        for h in headings:
            key = f"h{h.level}"
            level_counts[key] = level_counts.get(key, 0) + 1

        return {
            "total_paragraphs": total_paras,
            "text_paragraphs": text_paras,
            "word_count": word_count,
            "character_count": char_count,
            "heading_count": len(headings),
            "heading_levels": level_counts,
            "h1_count": h1,
            "h2_count": h2,
            "h3_plus_count": h3,
            "style_count": len(styles),
            "custom_styles": sum(1 for s in styles.values() if s.is_custom),
            "heading_styles": sum(1 for s in styles.values() if s.is_heading),
            "table_count": len(tables),
            "image_count": len(images),
            "reference_count": len(refs),
            "citation_count": len(citations),
            "footnote_count": len(footnotes or []),
            "endnote_count": sum(1 for f in (footnotes or []) if "endnote" in str(type(f)).lower()),
            "header_footer_count": len(headers_footers or []),
            "cross_reference_count": len(cross_references or []),
            "watermark_count": len(watermarks or []),
            "equation_count": len(equations or []),
            "section_types": self._count_section_types(sections),
        }

    def _count_section_types(self, sections: List[SectionInfo]) -> Dict[str, int]:
        counts: Dict[str, int] = {}
        for sec in sections:
            st = sec.section_type
            counts[st] = counts.get(st, 0) + 1
            child_counts = self._count_section_types(sec.children)
            for k, v in child_counts.items():
                counts[k] = counts.get(k, 0) + v
        return counts
