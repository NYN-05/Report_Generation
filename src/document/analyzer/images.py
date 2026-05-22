import re
from typing import Dict, List, Optional
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .models import ImageInfo


EMUS_PER_INCH = 914400


FIGURE_CAPTION_PATTERN = re.compile(
    r'^(figure|fig\.)\s+\d+(\.\d+)*[\.:\s]', re.IGNORECASE
)


class ImageDetector:
    """Detects and extracts metadata from all images in a DOCX document."""

    def __init__(self, doc: DocxDocument):
        self.doc = doc
        self._images: List[ImageInfo] = []
        self._heading_names: List[str] = []

    def detect(self, headings: Optional[List] = None) -> List[ImageInfo]:
        self._images = []
        if headings:
            self._heading_names = [h.text for h in headings]

        rels = self.doc.part.rels
        img_idx = 0

        for para_idx, para in enumerate(self.doc.paragraphs):
            images_in_para = self._extract_images_from_para(para, para_idx, img_idx)
            for img_info in images_in_para:
                self._images.append(img_info)
                img_idx += 1

        for table_idx, table in enumerate(self.doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        images_in_cell = self._extract_images_from_para(
                            para, -1, img_idx
                        )
                        for img_info in images_in_cell:
                            img_info.anchor_section = f"table_{table_idx}_row_{ri}"
                            self._images.append(img_info)
                            img_idx += 1

        self._attach_captions()
        self._resolve_anchor_sections()

        return self._images

    def _extract_images_from_para(self, para, para_idx: int, start_idx: int) -> List[ImageInfo]:
        images: List[ImageInfo] = []
        elem = para._element

        for drawing in elem.iter(qn('w:drawing')):
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is None:
                continue
            rId = blip.get(qn('r:embed') or qn('r:link'))
            if rId is None:
                continue

            img_info = ImageInfo(
                index=start_idx + len(images),
                rId=rId,
                paragraph_index=para_idx,
            )

            extent = drawing.find('.//' + qn('wp:extent'))
            if extent is not None:
                cx = extent.get('cx')
                cy = extent.get('cy')
                if cx:
                    img_info.width_emus = int(cx)
                    img_info.width_inches = int(cx) / EMUS_PER_INCH
                if cy:
                    img_info.height_emus = int(cy)
                    img_info.height_inches = int(cy) / EMUS_PER_INCH

            docPr = drawing.find('.//' + qn('wp:docPr'))
            if docPr is not None:
                img_info.alt_text = docPr.get('descr', '')

            images.append(img_info)

        return images

    def _attach_captions(self):
        MAX_CAPTION_DISTANCE = 20
        caption_map: Dict[int, str] = {}
        for pi, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if FIGURE_CAPTION_PATTERN.match(text):
                caption_map[pi] = text

        used_captions: set = set()
        for img in self._images:
            best_caption = None
            best_score = float('inf')
            best_cpi = -1

            for cpi, caption_text in caption_map.items():
                if cpi in used_captions:
                    continue
                distance = abs(img.paragraph_index - cpi)
                if distance > MAX_CAPTION_DISTANCE:
                    continue
                score = distance - 0.5 if cpi > img.paragraph_index else distance

                if score < best_score:
                    best_score = score
                    best_caption = caption_text
                    best_cpi = cpi

            if best_caption is not None:
                img.caption = best_caption
                if best_cpi >= 0:
                    used_captions.add(best_cpi)

    def _resolve_anchor_sections(self):
        for img in self._images:
            if img.paragraph_index < 0:
                continue
            best_heading = None
            for h in self._heading_names:
                for pi in range(img.paragraph_index - 1, -1, -1):
                    if pi < len(self.doc.paragraphs):
                        if self.doc.paragraphs[pi].text.strip() == h:
                            best_heading = h
                            break
                if best_heading:
                    break
            if best_heading:
                img.anchor_section = best_heading

    def get_image_count(self) -> int:
        return len(self._images)
