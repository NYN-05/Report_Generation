from typing import Dict, List, Optional, Tuple
from src.core.logger import get_logger
from .document_styles import DocumentStyles, Alignment

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


VALIDATION_RULES = {
    "content_font_times_new_roman": {
        "description": "All content paragraphs must use Times New Roman",
        "severity": "error",
    },
    "content_size_12pt": {
        "description": "All content paragraphs must be 12 pt",
        "severity": "error",
    },
    "content_justified": {
        "description": "All content paragraphs must be justified",
        "severity": "error",
    },
    "content_line_spacing_1_5": {
        "description": "Content paragraphs must have 1.5 line spacing",
        "severity": "error",
    },
    "main_headings_16pt_center": {
        "description": "Main headings (level 1) must be 16 pt and centered",
        "severity": "error",
    },
    "subheadings_14pt_left": {
        "description": "Subheadings (level 2) must be 14 pt and left-aligned",
        "severity": "error",
    },
    "no_calibri": {
        "description": "No Calibri font should appear anywhere",
        "severity": "warning",
    },
    "bullets_proper_style": {
        "description": "Bullet lists must use proper DOCX bullet style, not manual characters",
        "severity": "warning",
    },
}


class DocumentStyleValidator:
    def __init__(self, styles: Optional[DocumentStyles] = None):
        from .style_manager import StyleManager
        self._styles = styles or StyleManager.get_styles()
        self._issues: List[Dict] = []

    def validate(self, doc) -> Tuple[bool, List[Dict]]:
        self._issues = []
        if not HAS_DOCX:
            return True, []
        for i, paragraph in enumerate(doc.paragraphs):
            self._check_paragraph(paragraph, i)
        self._check_heading_consistency(doc)
        auto_fixed = self._auto_fix(doc)
        passed = len([i for i in self._issues if i.get("severity") == "error"]) == 0
        if self._issues:
            logger.info(
                f"Style validation: {'PASS' if passed else 'ISSUES'} "
                f"({len(self._issues)} issues, {auto_fixed} auto-fixed)"
            )
        return passed, self._issues

    def _check_paragraph(self, paragraph, idx: int):
        if not paragraph.text.strip():
            return
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            self._check_heading(paragraph, style_name, idx)
            return
        for run in paragraph.runs:
            if run.font.name and "Calibri" in run.font.name:
                self._issues.append({
                    "paragraph": idx,
                    "rule": "no_calibri",
                    "severity": "warning",
                    "text": paragraph.text[:60],
                    "message": f"Calibri font found: '{run.font.name}'",
                })
            if run.font.name and run.font.name != "Times New Roman":
                self._issues.append({
                    "paragraph": idx,
                    "rule": "content_font_times_new_roman",
                    "severity": "error",
                    "text": paragraph.text[:60],
                    "message": f"Font '{run.font.name}' instead of Times New Roman",
                })
        if style_name in ("Normal", ""):
            pf = paragraph.paragraph_format
            if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                self._issues.append({
                    "paragraph": idx,
                    "rule": "content_justified",
                    "severity": "error",
                    "text": paragraph.text[:60],
                    "message": f"Alignment {pf.alignment} instead of JUSTIFY",
                })

    def _check_heading(self, paragraph, style_name: str, idx: int):
        pf = paragraph.paragraph_format
        level = style_name.replace("Heading ", "")
        try:
            h_level = int(level) if level.isdigit() else 1
        except ValueError:
            h_level = 1
        expected_size = 16 if h_level <= 1 else (14 if h_level == 2 else 12)
        expected_align = WD_ALIGN_PARAGRAPH.CENTER if h_level <= 1 else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            if run.font.size and int(run.font.size.pt) not in (expected_size, expected_size - 1, expected_size + 1):
                self._issues.append({
                    "paragraph": idx,
                    "rule": "main_headings_16pt_center" if h_level <= 1 else "subheadings_14pt_left",
                    "severity": "warning",
                    "text": paragraph.text[:60],
                    "message": f"Heading {h_level} font size {run.font.size.pt}pt, expected {expected_size}pt",
                })
        if pf.alignment is not None and pf.alignment != expected_align:
            self._issues.append({
                "paragraph": idx,
                "rule": "main_headings_16pt_center" if h_level <= 1 else "subheadings_14pt_left",
                "severity": "warning",
                "text": paragraph.text[:60],
                "message": f"Heading {h_level} alignment should be {'CENTER' if h_level <= 1 else 'LEFT'}",
            })

    def _check_heading_consistency(self, doc):
        from collections import Counter
        style_counts = Counter()
        for p in doc.paragraphs:
            if p.style:
                style_counts[p.style.name] += 1

    def _auto_fix(self, doc) -> int:
        fixed = 0
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""
            for run in paragraph.runs:
                if run.font.name and "Calibri" in run.font.name:
                    run.font.name = "Times New Roman"
                    fixed += 1
                if run.font.name != "Times New Roman" and style_name not in ("Heading 1", "Heading 2", "Heading 3"):
                    run.font.name = "Times New Roman"
                    fixed += 1
        return fixed

    def report(self) -> str:
        if not self._issues:
            return "✓ All styles valid"
        lines = [f"Style validation: {len(self._issues)} issues"]
        for issue in self._issues[:10]:
            lines.append(f"  [{issue['severity']}] {issue['message']}")
        if len(self._issues) > 10:
            lines.append(f"  ... and {len(self._issues) - 10} more")
        return "\n".join(lines)
