"""
Style Analyzer Module
====================
Analyzes styles from DOCX templates.
"""

import os
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from src.core.logger import get_logger
from src.core.exceptions import StyleExtractionError

logger = get_logger(__name__)


class StyleAnalyzer:
    """Analyzes and extracts styles from DOCX documents."""

    def __init__(self, document_path: str):
        self.document_path = document_path
        self._document: Optional[Document] = None

    def analyze(self) -> Dict[str, Any]:
        """Analyze all styles in the document."""
        if not os.path.exists(self.document_path):
            raise StyleExtractionError(f"Document not found: {self.document_path}")

        try:
            self._document = Document(self.document_path)
            return {
                "character_styles": self._get_character_styles(),
                "paragraph_styles": self._get_paragraph_styles(),
                "table_styles": self._get_table_styles(),
                "default_font": self._get_default_font(),
                "default_paragraph": self._get_default_paragraph(),
            }
        except Exception as e:
            raise StyleExtractionError(f"Style analysis failed: {e}")

    def _extract_style_properties(self, style_elem) -> Dict[str, Any]:
        """Extract full font and paragraph properties from a style XML element."""
        props = {}
        style_id = style_elem.get(qn('w:styleId'))
        props['name'] = style_id

        name_elem = style_elem.find(qn('w:name'))
        if name_elem is not None:
            props['display_name'] = name_elem.get(qn('w:val'))

        rPr = style_elem.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                props['font_name'] = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))

            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                half_pts = sz.get(qn('w:val'))
                if half_pts:
                    props['font_size'] = int(half_pts) / 2

            szCs = rPr.find(qn('w:szCs'))
            if szCs is not None:
                half_pts = szCs.get(qn('w:val'))
                if half_pts:
                    props['font_size_complex'] = int(half_pts) / 2

            b = rPr.find(qn('w:b'))
            if b is not None:
                val = b.get(qn('w:val'))
                props['bold'] = val is None or val.lower() in ('1', 'true', 'on')

            i = rPr.find(qn('w:i'))
            if i is not None:
                val = i.get(qn('w:val'))
                props['italic'] = val is None or val.lower() in ('1', 'true', 'on')

            u = rPr.find(qn('w:u'))
            if u is not None:
                props['underline'] = u.get(qn('w:val')) or 'single'

            color = rPr.find(qn('w:color'))
            if color is not None:
                props['color'] = color.get(qn('w:val'))

        pPr = style_elem.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                props['alignment'] = jc.get(qn('w:val'))

            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                after = spacing.get(qn('w:after'))
                if after is not None:
                    props['space_after'] = int(after) / 20
                before = spacing.get(qn('w:before'))
                if before is not None:
                    props['space_before'] = int(before) / 20
                line = spacing.get(qn('w:line'))
                if line is not None:
                    props['line_spacing'] = int(line) / 240
                line_rule = spacing.get(qn('w:lineRule'))
                if line_rule is not None:
                    props['line_spacing_rule'] = line_rule

            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                left = ind.get(qn('w:left'))
                if left is not None:
                    props['indent_left'] = int(left) / 20
                right = ind.get(qn('w:right'))
                if right is not None:
                    props['indent_right'] = int(right) / 20
                first_line = ind.get(qn('w:firstLine'))
                if first_line is not None:
                    props['first_line_indent'] = int(first_line) / 20
                hanging = ind.get(qn('w:hanging'))
                if hanging is not None:
                    props['hanging_indent'] = int(hanging) / 20

        return props

    def _get_character_styles(self) -> List[Dict[str, Any]]:
        """Get all character styles with full properties."""
        styles = []

        for style in self._document.styles.element.iter():
            if style.tag.endswith('rStyle'):
                style_id = style.get(qn('w:styleId'))
                if style_id:
                    props = self._extract_style_properties(style)
                    props['type'] = 'character'
                    styles.append(props)

        return styles

    def _get_paragraph_styles(self) -> List[Dict[str, Any]]:
        """Get all paragraph styles with full properties."""
        styles = []

        for style in self._document.styles.element.iter():
            if style.tag.endswith('pStyle'):
                style_id = style.get(qn('w:styleId'))
                if style_id:
                    props = self._extract_style_properties(style)
                    props['type'] = 'paragraph'
                    styles.append(props)

        return styles

    def _get_table_styles(self) -> List[Dict[str, Any]]:
        """Get all table styles."""
        styles = []

        for style in self._document.styles.element.iter():
            if style.tag.endswith('tblStyle'):
                style_name = style.get(qn('w:val'))
                if style_name:
                    styles.append({
                        "name": style_name,
                        "type": "table"
                    })

        return styles

    def _get_default_font(self) -> Dict[str, Any]:
        """Get default font settings from document styles."""
        try:
            default_rpr = self._document.styles.element.xpath(
                './/w:rPrDefault//w:rPr'
            )[0]

            font_name = "Calibri"
            font_size = 11

            rFonts = default_rpr.find(qn('w:rFonts'))
            if rFonts is not None:
                font_name = rFonts.get(qn('w:ascii')) or font_name

            sz = default_rpr.find(qn('w:sz'))
            if sz is not None:
                half_pts = sz.get(qn('w:val'))
                if half_pts:
                    font_size = int(half_pts) / 2

            return {"name": font_name, "size": font_size}
        except Exception:
            return {"name": "Calibri", "size": 11}

    def _get_default_paragraph(self) -> Dict[str, Any]:
        """Get default paragraph settings."""
        try:
            default_ppr = self._document.styles.element.xpath(
                './/w:pPrDefault//w:pPr'
            )
            if default_ppr:
                ppr = default_ppr[0]
                props = {}
                jc = ppr.find(qn('w:jc'))
                if jc is not None:
                    props['alignment'] = jc.get(qn('w:val'))
                spacing = ppr.find(qn('w:spacing'))
                if spacing is not None:
                    after = spacing.get(qn('w:after'))
                    if after is not None:
                        props['space_after'] = int(after) / 20
                    before = spacing.get(qn('w:before'))
                    if before is not None:
                        props['space_before'] = int(before) / 20
                    line = spacing.get(qn('w:line'))
                    if line is not None:
                        props['line_spacing'] = int(line) / 240
                return props
        except Exception:
            pass
        return {
            "alignment": "left",
            "space_after": 0,
            "line_spacing": 1.15
        }

    def extract_heading_styles(self) -> Dict[str, Dict]:
        """Extract all heading styles with their properties."""
        heading_styles = {}

        for para in self._document.paragraphs:
            style_name = para.style.name
            if style_name.startswith('Heading'):
                heading_styles[style_name] = {
                    "name": style_name,
                    "text": para.text[:50]
                }

        return heading_styles

    def get_style_mapping(self) -> Dict[str, str]:
        """Get mapping of style names to their properties."""
        mapping = {}

        for para in self._document.paragraphs:
            if para.style and para.style.name not in mapping:
                mapping[para.style.name] = {
                    "name": para.style.name,
                    "type": str(para.style.type) if para.style.type else "unknown"
                }

        return mapping


def analyze_styles(path: str) -> Dict[str, Any]:
    """Convenience function to analyze styles."""
    analyzer = StyleAnalyzer(path)
    return analyzer.analyze()