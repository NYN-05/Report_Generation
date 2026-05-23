import os
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import BaseDocument, DocumentMetadata, Section
from src.core.logger import get_logger
from src.document.styles import StyleManager

logger = get_logger(__name__)


class DocumentBuilder(BaseDocument):
    def __init__(self, metadata: Optional[DocumentMetadata] = None):
        super().__init__(metadata)
        self._document: Optional[Document] = None
        self._styles = StyleManager.get_instance()

    def create(self) -> "DocumentBuilder":
        self._document = Document()
        self._styles.setup_document(self._document)
        return self

    def add_cover_page(self, title: str, subtitle: str = "",
                       author: str = "", date: str = "") -> "DocumentBuilder":
        if not self._document:
            self.create()
        s = self._styles.get_styles()
        title_p = self._document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._styles.write_run(title_p, title, s.cover_page.title_font)
        if subtitle:
            sub_p = self._document.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(sub_p, subtitle, s.cover_page.subtitle_font)
        self._document.add_paragraph()
        if author or date:
            info_p = self._document.add_paragraph()
            info_parts = []
            if author:
                info_parts.append(f"Author: {author}")
            if date:
                info_parts.append(f"Date: {date}")
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(info_p, "\n".join(info_parts), s.cover_page.author_font)
        self._document.add_page_break()
        return self

    def add_toc(self, entries: List[str]) -> "DocumentBuilder":
        if not self._document:
            self.create()
        self._document.add_heading("Table of Contents", level=1)
        for i, entry in enumerate(entries, 1):
            p = self._document.add_paragraph(f"{i}. {entry}")
            p.paragraph_format.space_before = Pt(6)
        self._document.add_page_break()
        return self

    def add_section(self, title: str, content: str = "",
                    level: int = 1, add_page_break: bool = False) -> "DocumentBuilder":
        if not self._document:
            self.create()
        self._document.add_heading(title, level=level)
        if content:
            if isinstance(content, list):
                content = "\n\n".join(str(item) for item in content)
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    p = self._document.add_paragraph(para_text)
                    self._styles.apply_paragraph_style(p, self._styles.get_styles().content)
        if add_page_break:
            self._document.add_page_break()
        return self

    def add_paragraph(self, text: str, style: str = "Normal",
                      bold: bool = False, italic: bool = False) -> "DocumentBuilder":
        if not self._document:
            self.create()
        p = self._document.add_paragraph(text)
        if style != "Normal":
            p.style = style
        for run in p.runs:
            run.font.bold = bold
            run.font.italic = italic
            run.font.name = self._styles.get_styles().content.font.name
            run.font.size = Pt(self._styles.get_styles().content.font.size)
        return self

    def add_bullet_list(self, items: List[str]) -> "DocumentBuilder":
        if not self._document:
            self.create()
        for item in items:
            p = self._document.add_paragraph(item, style="List Bullet")
        return self

    def add_numbered_list(self, items: List[str]) -> "DocumentBuilder":
        if not self._document:
            self.create()
        for item in items:
            p = self._document.add_paragraph(item, style="List Number")
        return self

    def add_table(self, data: List[List[str]], headers: Optional[List[str]] = None,
                  style: str = "Table Grid") -> "DocumentBuilder":
        if not self._document:
            self.create()
        rows = len(data)
        cols = len(data[0]) if data else 0
        if headers:
            rows += 1
        table = self._document.add_table(rows=rows, cols=cols)
        table.style = style
        row_idx = 0
        if headers:
            for col_idx, header in enumerate(headers):
                table.rows[row_idx].cells[col_idx].text = header
            row_idx += 1
        for row in data:
            for col_idx, cell in enumerate(row):
                table.rows[row_idx].cells[col_idx].text = cell
            row_idx += 1
        return self

    def save(self, path: str) -> bool:
        if not self._document:
            logger.error("No document to save")
            return False
        try:
            from src.document.styles import DocumentStyleValidator
            validator = DocumentStyleValidator()
            validator.validate(self._document)
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            self._document.save(path)
            self.metadata.modified_at = datetime.now()
            self.metadata.paragraph_count = len(self._document.paragraphs)
            self.metadata.table_count = len(self._document.tables)
            logger.info(f"Document saved: {path}")
            return True
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return False

    def load(self, path: str) -> bool:
        try:
            self._document = Document(path)
            logger.info(f"Document loaded: {path}")
            return True
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return False

    def get_sections(self) -> List[Section]:
        if not self._document:
            return []
        sections = []
        for para in self._document.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                sections.append(Section(title=para.text, content="", level=level))
        return sections
