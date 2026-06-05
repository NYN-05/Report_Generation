import os
import re
from typing import Dict, Any, Optional, List
from src.core.logger import get_logger
from src.generator.content_blocks import (
    SectionContent, ParagraphBlock, BulletListBlock, BulletItem,
    TableBlock, TableRow, FigureBlock, HeadingBlock,
    SourceRequiredBlock, ContentBlock,
)
from docx.shared import RGBColor
from src.document.styles import StyleManager, DocumentStyleValidator

logger = get_logger(__name__)
_CONTROL_CHARS = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]')
_MAX_TEXT_LENGTH = 50000


def sanitize_text(text: str, max_length: int = _MAX_TEXT_LENGTH) -> str:
    if not text:
        return ""
    text = _CONTROL_CHARS.sub("", text)
    text = text[:max_length]
    return text

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXV2Generator:

    def __init__(self):
        self._doc: Optional[Document] = None
        self._styles = StyleManager.get_instance()

    def generate(
        self,
        title: str,
        author: str = "",
        subtitle: str = "",
        metadata: Optional[Dict] = None,
        sections: Optional[List[SectionContent]] = None,
        output_path: str = "output.docx",
        validate: bool = True,
        executive_summary: str = "",
        key_findings_section: Optional[SectionContent] = None,
        utilization_summary: Optional[Dict] = None,
    ) -> str:
        self._doc = Document()
        self._styles.setup_document(self._doc)
        meta = metadata or {}
        self._add_cover_page(title, author, subtitle, meta)
        self._add_table_of_contents(sections or [], has_exec=bool(executive_summary),
                                     has_kf=key_findings_section is not None)
        if executive_summary:
            self._add_executive_summary(executive_summary)
        if key_findings_section:
            self._render_section(key_findings_section)
        for section in sections or []:
            self._render_section(section)
        if utilization_summary:
            self._add_utilization_summary(utilization_summary)
        self._add_evidence_appendix(sections or [])
        self._add_references(sections or [])
        if validate and HAS_DOCX:
            validator = DocumentStyleValidator()
            passed, issues = validator.validate(self._doc)
            if issues:
                for issue in issues[:5]:
                    logger.info(f"  [{issue['severity']}] {issue['message']}")
        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self._doc.save(output_path)
        size_kb = os.path.getsize(output_path) / 1024 if os.path.exists(output_path) else 0
        logger.info(f"DOCX saved: {output_path} ({size_kb:.1f} KB, {len(sections or [])} sections)")
        return output_path

    def _add_cover_page(self, title: str, author: str,
                         subtitle: str = "",
                         metadata: Optional[Dict] = None):
        meta = metadata or {}
        for _ in range(6):
            self._doc.add_paragraph()
        title_p = self._doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._styles.write_run(title_p, title, self._styles.get_styles().cover_page.title_font)
        if subtitle:
            self._doc.add_paragraph()
            sub_p = self._doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(sub_p, subtitle, self._styles.get_styles().cover_page.author_font)
        if author:
            self._doc.add_paragraph()
            auth_p = self._doc.add_paragraph()
            auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(auth_p, f"Author: {author}", self._styles.get_styles().cover_page.author_font)
        if meta.get("domain") or meta.get("report_type"):
            self._doc.add_paragraph()
            info_p = self._doc.add_paragraph()
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parts = []
            if meta.get("domain"):
                parts.append(f"Domain: {meta['domain']}")
            if meta.get("report_type"):
                rtype = meta["report_type"].replace("_", " ").title()
                parts.append(f"Type: {rtype}")
            if meta.get("audience"):
                parts.append(f"Audience: {meta['audience']}")
            if parts:
                self._styles.write_run(info_p, " | ".join(parts), self._styles.get_styles().cover_page.author_font)
        from datetime import date
        self._doc.add_paragraph()
        date_p = self._doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._styles.write_run(date_p, date.today().strftime("%B %d, %Y"), self._styles.get_styles().cover_page.author_font)
        self._doc.add_page_break()

    def _add_table_of_contents(self, sections: List[SectionContent],
                                 has_exec: bool = False,
                                 has_kf: bool = False):
        toc_heading = self._doc.add_heading("Table of Contents", level=1)
        i = 1
        if has_exec:
            p = self._doc.add_paragraph(f"{i}. Executive Summary")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            i += 1
        if has_kf:
            p = self._doc.add_paragraph(f"{i}. Key Findings and Insights")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            i += 1
        for section in sections:
            p = self._doc.add_paragraph(f"{i}. {section.heading}")
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(2)
            i += 1
        p = self._doc.add_paragraph(f"{i}. Evidence Utilization Summary")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        i += 1
        p = self._doc.add_paragraph(f"{i}. Appendix: Evidence Traceability")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        i += 1
        p = self._doc.add_paragraph(f"{i}. References")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        self._doc.add_page_break()

    def _add_executive_summary(self, summary: str):
        heading = self._doc.add_heading("Executive Summary", level=1)
        p = self._doc.add_paragraph()
        self._styles.apply_paragraph_style(p, self._styles.get_styles().content)
        run = p.add_run(sanitize_text(summary))
        run.font.name = self._styles.get_styles().content.font.name
        run.font.size = Pt(self._styles.get_styles().content.font.size)
        self._doc.add_page_break()

    def _add_utilization_summary(self, summary: Dict):
        self._doc.add_page_break()
        heading = self._doc.add_heading("Evidence Utilization Summary", level=1)

        facts_table = self._doc.add_table(rows=6, cols=2)
        facts_table.style = "Light Shading Accent 1"
        facts_table.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def set_cell(row, col, text, bold=False):
            cell = facts_table.rows[row].cells[col]
            cell.text = str(text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.bold = bold

        set_cell(0, 0, "Metric", True)
        set_cell(0, 1, "Value", True)
        set_cell(1, 0, "Total Facts Collected")
        set_cell(1, 1, summary.get("total_facts", 0))
        set_cell(2, 0, "Facts Assigned to Clusters")
        set_cell(2, 1, summary.get("assigned_facts", 0))
        set_cell(3, 0, "Facts Incorporated in Report")
        set_cell(3, 1, summary.get("generated_facts", 0))
        set_cell(4, 0, "Fact Utilization Rate")
        set_cell(4, 1, f"{summary.get('utilization_rate', 0) * 100:.1f}%")
        set_cell(5, 0, "Source Coverage Rate")
        set_cell(5, 1, f"{summary.get('source_coverage_rate', 0) * 100:.0f}%")

        self._doc.add_paragraph()
        p = self._doc.add_paragraph()
        run = p.add_run(
            f"Sources: {summary.get('covered_sources', 0)} of "
            f"{summary.get('total_sources', 0)} used | "
            f"Knowledge areas: {summary.get('total_clusters', 0)}"
        )
        run.font.size = Pt(10)
        run.font.italic = True

        unused = summary.get("unused_by_category", {})
        if unused and any(v > 0 for v in unused.values()):
            self._doc.add_paragraph()
            self._doc.add_heading("Unused Fact Profile", level=2)
            for cat, count in unused.items():
                if count > 0:
                    label = cat.replace("_", " ").title()
                    p = self._doc.add_paragraph(f"  \u2022 {label}: {count}")
                    p.paragraph_format.space_after = Pt(2)

    def _add_evidence_appendix(self, sections: List[SectionContent]):
        appendix_heading = self._doc.add_heading("Appendix: Evidence Traceability", level=1)
        for section in sections:
            self._doc.add_heading(section.heading, level=2)
            for block in section.blocks:
                if isinstance(block, ParagraphBlock) and block.text:
                    p = self._doc.add_paragraph()
                    run = p.add_run(block.text[:200] + ("..." if len(block.text) > 200 else ""))
                    run.font.size = Pt(10)
                    run.italic = True
                    if block.evidence_source:
                        src_run = p.add_run(f"  [{block.evidence_source}]")
                        src_run.font.size = Pt(9)
                        src_run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_references(self, sections: List[SectionContent]):
        heading = self._doc.add_heading("References", level=1)
        seen_sources: Dict[str, int] = {}
        for section in sections:
            for block in section.blocks:
                if isinstance(block, ParagraphBlock):
                    for c in block.citations:
                        if c.source and c.source not in seen_sources:
                            seen_sources[c.source] = len(seen_sources) + 1
                elif isinstance(block, BulletListBlock):
                    for item in block.items:
                        for c in item.citations:
                            if c.source and c.source not in seen_sources:
                                seen_sources[c.source] = len(seen_sources) + 1

        if not seen_sources:
            p = self._doc.add_paragraph("No external references were cited in this report.")
            p.runs[0].font.italic = True
            return

        for source, idx in seen_sources.items():
            p = self._doc.add_paragraph(f"[{idx}] {source}")
            p.paragraph_format.space_after = Pt(2)

    def _render_section(self, section: SectionContent):
        for block in section.blocks:
            self._render_block(block)

    def _render_block(self, block: ContentBlock):
        if isinstance(block, HeadingBlock):
            self._render_heading(block)
        elif isinstance(block, ParagraphBlock):
            self._render_paragraph(block)
        elif isinstance(block, BulletListBlock):
            self._render_bullet_list(block)
        elif isinstance(block, TableBlock):
            self._render_table(block)
        elif isinstance(block, FigureBlock):
            self._render_figure(block)
        elif isinstance(block, SourceRequiredBlock):
            self._render_source_required(block)

    def _render_heading(self, block: HeadingBlock):
        level = min(block.level, 4)
        heading = self._doc.add_heading(block.text, level=level)
        hs = self._styles.get_styles().get_heading(level)
        if heading.runs:
            for run in heading.runs:
                run.font.name = hs.font.name
                run.font.size = Pt(hs.font.size)
                run.font.bold = hs.font.bold
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 1 else WD_ALIGN_PARAGRAPH.LEFT

    def _render_paragraph(self, block: ParagraphBlock):
        if not block.text.strip():
            return
        text = sanitize_text(block.text.strip())
        p = self._doc.add_paragraph()
        self._styles.apply_paragraph_style(p, self._styles.get_styles().content)
        run = p.add_run(text)
        run.font.name = self._styles.get_styles().content.font.name
        run.font.size = Pt(self._styles.get_styles().content.font.size)
        if block.citations:
            cite_text = " ".join(
                f"[{c.source}]" for c in block.citations if c.source
            )
            if cite_text:
                cite_run = p.add_run(f" {cite_text}")
                cite_run.font.size = Pt(10)
                cite_run.font.color.rgb = RGBColor(*self._styles._parse_color("#808080"))

    def _render_bullet_list(self, block: BulletListBlock):
        s = self._styles.get_styles()
        for item in block.items:
            bp = self._doc.add_paragraph(style="List Bullet")
            self._styles.apply_paragraph_style(bp, s.bullet)
            bp.clear()
            run = bp.add_run(f"{item.title}: " if item.title else "")
            run.font.name = s.bullet.font.name
            run.font.size = Pt(s.bullet.font.size)
            run.bold = True
            if item.description:
                drun = bp.add_run(item.description)
                drun.font.name = s.bullet.font.name
                drun.font.size = Pt(s.bullet.font.size)
            if item.citations:
                cite = " ".join(f"[{c.source}]" for c in item.citations if c.source)
                if cite:
                    crun = bp.add_run(f" {cite}")
                    crun.font.size = Pt(10)

    def _render_table(self, block: TableBlock):
        s = self._styles.get_styles()
        if block.caption:
            cap_p = self._doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._styles.write_run(cap_p, f"Table: {block.caption}", s.table.title_font)
            cap_p.paragraph_format.space_after = Pt(4)
        rows_data = [block.headers] + [r.cells for r in block.rows]
        if not rows_data or not rows_data[0]:
            return
        num_cols = len(rows_data[0])
        table = self._doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = s.table.grid_style
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_text)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(s.table.cell_font.size)
                        if i == 0:
                            run.font.bold = True
                            run.font.size = Pt(s.table.header_font.size)

    def _render_figure(self, block: FigureBlock):
        s = self._styles.get_styles()
        fp = self._doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._styles.write_run(fp, f"[Figure: {block.caption}]", s.figure.caption_font)
        if block.description:
            dp = self._doc.add_paragraph(block.description)
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in dp.runs:
                r.font.size = Pt(s.figure.caption_font.size)
                r.font.italic = True

    def _render_source_required(self, block: SourceRequiredBlock):
        p = self._doc.add_paragraph()
        run = p.add_run(block.message)
        run.italic = True
        run.font.color.rgb = RGBColor(*self._styles._parse_color("#B43232"))
        run.font.size = Pt(10)
        run.font.name = self._styles.get_styles().content.font.name
