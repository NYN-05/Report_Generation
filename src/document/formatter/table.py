"""
Table Formatter Module
======================
Table formatting utilities.
"""

from typing import List, Optional
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from src.core.logger import get_logger

logger = get_logger(__name__)


class TableFormatter:
    """Formats table properties."""

    @staticmethod
    def format_table(
        table,
        column_widths: List[int] = None,
        alignment: WD_TABLE_ALIGNMENT = None
    ):
        """Apply table formatting."""
        if alignment:
            table.alignment = alignment

        if column_widths:
            TableFormatter.set_column_widths(table, column_widths)

    @staticmethod
    def set_column_widths(table, widths: List[int]):
        """Set column widths."""
        for idx, width in enumerate(widths):
            if idx < len(table.columns):
                for cell in table.columns[idx].cells:
                    cell.width = Inches(width / 72)

    @staticmethod
    def format_header_row(row):
        """Format header row with bold and background."""
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

            tc = cell._element
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = tc.makeelement(qn('w:tcPr'))
                tc.append(tc_pr)

            shd = tc_pr.find(qn('w:shd'))
            if shd is None:
                shd = tc_pr.makeelement(qn('w:shd'))
                tc_pr.append(shd)
                shd.set(qn('w:fill'), 'D9D9D9')

    @staticmethod
    def format_cell(cell, bold: bool = False, font_size: int = 11):
        """Format table cell."""
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = bold
                if font_size:
                    run.font.size = Pt(font_size)

    @staticmethod
    def apply_borders(table, color: RGBColor = None):
        """Apply borders to table."""
        if color is None:
            color = RGBColor(200, 200, 200)

        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tc_pr = tc.find(qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = tc.makeelement(qn('w:tcPr'))
                    tc.append(tc_pr)

                tc_borders = tc_pr.find(qn('w:tcBorders'))
                if tc_borders is None:
                    tc_borders = tc_pr.makeelement(qn('w:tcBorders'))
                    tc_pr.append(tc_borders)

    @staticmethod
    def create_styled_table(rows: int, columns: int, headers: List[str] = None):
        """Create a styled table with headers."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=rows + (1 if headers else 0), cols=columns)
        table.style = 'Table Grid'

        if headers:
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            TableFormatter.format_header_row(table.rows[0])

        return table