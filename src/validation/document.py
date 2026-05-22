"""
Document Validator Module
========================
Validates document structure.
"""

import os
from typing import Dict, Any
from .base import Validator, ValidationResult
from src.core.logger import get_logger

logger = get_logger(__name__)


class DocumentValidator(Validator):
    """Validates document structure and content."""

    def validate(self, document_path: str) -> ValidationResult:
        """Validate a document file."""
        self._reset()

        if not os.path.exists(document_path):
            self.add_error("file", f"Document not found: {document_path}")
            return self._create_result()

        if not document_path.endswith('.docx'):
            self.add_error("format", "Document must be a .docx file")

        try:
            from docx import Document
            doc = Document(document_path)

            if len(doc.paragraphs) == 0:
                self.add_warning("content", "Document has no paragraphs")

            if len(doc.sections) == 0:
                self.add_warning("structure", "Document has no sections")

            for idx, para in enumerate(doc.paragraphs):
                if len(para.text) > 10000:
                    self.add_warning(f"paragraph_{idx}", "Paragraph is very long")

        except Exception as e:
            self.add_error("parsing", f"Failed to parse document: {e}")

        return self._create_result()

    def validate_structure(self, content: Dict[str, Any]) -> ValidationResult:
        """Validate content structure."""
        self._reset()

        required_sections = ['title', 'introduction']
        for section in required_sections:
            if section not in content or not content[section]:
                self.add_error(section, f"Required section '{section}' is missing")

        if 'sections' in content:
            if not isinstance(content['sections'], list):
                self.add_error("sections", "Sections must be an array")
            elif len(content['sections']) == 0:
                self.add_warning("sections", "No content sections defined")

        return self._create_result()

    def validate_metadata(self, metadata: Dict[str, Any]) -> ValidationResult:
        """Validate document metadata."""
        self._reset()

        if 'title' in metadata:
            if not isinstance(metadata['title'], str):
                self.add_error("title", "Title must be a string")
            elif len(metadata['title']) > 200:
                self.add_error("title", "Title too long (max 200 chars)")

        if 'author' in metadata:
            if not isinstance(metadata['author'], str):
                self.add_error("author", "Author must be a string")

        return self._create_result()


def validate_document(path: str) -> ValidationResult:
    """Convenience function to validate a document."""
    validator = DocumentValidator()
    return validator.validate(path)