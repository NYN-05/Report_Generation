"""
Export Agent
============
Handles DOCX/PDF/PDF-A document export.
"""

import os
from typing import Any, Dict, Optional
from .base import BaseAgent, AgentResponse
from src.core.logger import get_logger

logger = get_logger(__name__)


class ExportAgent(BaseAgent):
    """Agent responsible for document export in multiple formats.

    Responsibilities:
    - DOCX generation
    - PDF conversion (docx2pdf, LibreOffice, win32com)
    - PDF/A compliance
    - Format validation before export
    """

    def __init__(self, provider=None):
        super().__init__("export", provider)

    def execute(self, input_data: Any, **kwargs) -> AgentResponse:
        if not isinstance(input_data, dict):
            return self._create_response(False, error="Input must be a dict")

        plan = input_data.get("plan")
        output_path = input_data.get("output_path", "output/output.docx")
        formats = input_data.get("formats", ["docx"])
        builder = input_data.get("builder")

        if not plan:
            return self._create_response(False, error="No report plan provided")

        results = {}

        if "docx" in formats:
            docx_result = self._export_docx(plan, output_path, builder)
            results["docx"] = docx_result
            if docx_result.get("success") and "pdf" in formats:
                pdf_path = output_path.replace(".docx", ".pdf")
                pdf_result = self._export_pdf(docx_result.get("path", ""), pdf_path)
                results["pdf"] = pdf_result

        return self._create_response(
            all(r.get("success") for r in results.values()),
            data={"exports": results, "formats": formats},
        )

    def _export_docx(self, plan, output_path, builder=None) -> Dict:
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            if builder:
                success = builder.build(plan, output_path)
                if success:
                    return {"success": True, "path": output_path, "format": "docx"}
                return {"success": False, "error": "Builder failed"}

            try:
                from src.document.blueprint import BlueprintBuilder
                b = BlueprintBuilder()
                success = b.build(plan, output_path)
                if success:
                    return {"success": True, "path": output_path, "format": "docx"}
            except Exception as blueprint_err:
                logger.warning(f"BlueprintBuilder failed, using fallback: {blueprint_err}")

            return self._fallback_docx_export(plan, output_path)

        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return {"success": False, "error": str(e), "format": "docx"}

    def _export_pdf(self, docx_path: str, pdf_path: str) -> Dict:
        from src.pipeline.export.pdf import PDFExportPipeline
        pipeline = PDFExportPipeline()
        result = pipeline.execute(docx_path, output_path=pdf_path)
        return {
            "success": result.success,
            "path": pdf_path,
            "format": "pdf",
            "error": result.error,
        }

    def _fallback_docx_export(self, plan, output_path: str) -> Dict:
        try:
            report_content = []
            if hasattr(plan, "sections"):
                for sec in plan.sections:
                    if sec.content:
                        report_content.append(f"# {sec.heading}\n\n{sec.content}")
            elif isinstance(plan, dict):
                sections = plan.get("sections", [])
                for sec in sections:
                    report_content.append(f"# {sec.get('heading', '')}\n\n{sec.get('content', '')}")

            if report_content:
                from docx import Document
                from docx.shared import Inches, Pt
                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                for block in report_content:
                    lines = block.split("\n")
                    for line in lines:
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.strip():
                            doc.add_paragraph(line)
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                doc.save(output_path)
                return {"success": True, "path": output_path, "format": "docx"}
            return {"success": False, "error": "No content to export"}
        except Exception as e:
            return {"success": False, "error": str(e), "format": "docx"}

    def export_docx(self, plan, output_path: str, builder=None) -> bool:
        result = self._export_docx(plan, output_path, builder)
        return result.get("success", False)

    def export_pdf(self, docx_path: str, pdf_path: str) -> bool:
        result = self._export_pdf(docx_path, pdf_path)
        return result.get("success", False)
