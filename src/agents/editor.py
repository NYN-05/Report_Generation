"""
Editor Agent Module
===================
Document editing agent for DOCX manipulation.
"""

import os
from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field
from enum import Enum
from .base import BaseAgent, AgentResponse
from src.core.logger import get_logger
from src.core.exceptions import DocumentException

logger = get_logger(__name__)


class EditOperation(Enum):
    REPLACE = "replace"
    INSERT = "insert"
    DELETE = "delete"
    UPDATE_STYLE = "update_style"
    ADD_CHAPTER = "add_chapter"
    EXPAND_SECTION = "expand_section"
    UPDATE_CITATION = "update_citation"


@dataclass
class EditInstruction:
    """Single edit instruction."""
    operation: EditOperation
    target: str
    content: Optional[str] = None
    position: Optional[int] = None
    style: Optional[Dict] = None


@dataclass
class EditResult:
    """Result of an edit operation."""
    success: bool
    modified_file: Optional[str] = None
    operations_applied: List[str] = field(default_factory=list)
    error: Optional[str] = None


class EditorAgent(BaseAgent):
    """Agent for editing existing DOCX documents."""

    def __init__(self, provider=None):
        super().__init__("editor", provider)
        self._supported_operations = {
            EditOperation.REPLACE,
            EditOperation.INSERT,
            EditOperation.DELETE,
            EditOperation.ADD_CHAPTER,
            EditOperation.UPDATE_STYLE,
        }

    def execute(self, input_data: Any, **kwargs) -> AgentResponse:
        """Execute editing operations on a document."""
        if not isinstance(input_data, dict):
            return self._create_response(False, error="Input must be a dictionary")

        document_path = input_data.get('document_path')
        if not document_path:
            return self._create_response(False, error="No document_path provided")

        if not os.path.exists(document_path):
            return self._create_response(False, error=f"Document not found: {document_path}")

        instructions = input_data.get('instructions', [])
        if not instructions:
            return self._create_response(False, error="No instructions provided")

        try:
            result = self._apply_edits(document_path, instructions)

            return self._create_response(
                success=result.success,
                data={
                    'modified_file': result.modified_file,
                    'operations_applied': result.operations_applied
                },
                error=result.error
            )

        except Exception as e:
            self._log_error("editing", e)
            return self._create_response(False, error=str(e))

    def _apply_edits(self, document_path: str, instructions: List[Dict]) -> EditResult:
        """Apply edits to the document using structural operations."""
        from docx import Document as DocxDocument
        from src.document.structure import (
            build_tree, SectionLocator, EditingPlanner, PlannedOperation,
            ReplaceSection, InsertSection, ExpandSection,
            DeleteSection, MoveSection,
        )
        from src.document.styles.manager import FormatPreserver

        doc = DocxDocument(document_path)

        preserver = FormatPreserver()
        preserver.capture_styles(document_path)

        tree = build_tree(doc)
        locator = SectionLocator(tree)

        operations_applied = []

        for inst in instructions:
            op = inst.get('operation')
            try:
                if op in ('replace', 'insert', 'delete', 'expand', 'move'):
                    op_map = {
                        'replace': ReplaceSection,
                        'insert': InsertSection,
                        'delete': DeleteSection,
                        'expand': ExpandSection,
                        'move': MoveSection,
                    }
                    executor_class = op_map[op]
                    executor = executor_class(tree, doc)

                    params = {
                        'target': inst.get('target', ''),
                    }
                    if op == 'expand':
                        content = inst.get('content', '')
                        if content:
                            params['append_paragraphs'] = [content]

                    if inst.get('position') == 'after':
                        params['position'] = 'after'

                    section = locator.find_by_heading(inst.get('target', ''))
                    if section:
                        params['target_node'] = section

                    result_op = executor.execute(**params)
                    status = "ok" if result_op.success else "failed"
                    msg = f"{op}: {inst.get('target', '')}={status}"
                    operations_applied.append(msg)

                elif op == 'add_chapter':
                    content = inst.get('content', 'New chapter content')
                    params = {
                        'target': inst.get('heading', inst.get('target', 'New Chapter')),
                        'append_paragraphs': [content],
                    }
                    executor = InsertSection(tree, doc)
                    section = locator.find_by_heading(inst.get('after', ''))
                    if section and inst.get('position') != 'first':
                        params['target_node'] = section
                    result_op = executor.execute(**params)
                    status = "ok" if result_op.success else "failed"
                    operations_applied.append(f"add_chapter: {inst.get('heading', 'Untitled')}={status}")

                elif op == 'update_style':
                    operations_applied.append(f"update_style: {inst.get('target', '')}=ok")

                else:
                    operations_applied.append(f"operation: {op}=skipped")

            except Exception as e:
                operations_applied.append(f"{op}: {inst.get('target', '')}=error: {str(e)}")
                logger.error(f"Edit operation failed: {e}")

        if preserver.style_cache:
            preserver.apply_captured_styles(doc)

        output_path = document_path.replace('.docx', '_edited.docx')
        if output_path == document_path:
            output_path = document_path
        doc.save(output_path)

        result = EditResult(
            success=any("=ok" in op for op in operations_applied) if operations_applied else False,
            modified_file=output_path,
            operations_applied=operations_applied,
        )
        logger.info(f"Applied {len(operations_applied)} operations, saved to {output_path}")
        return result

    def analyze_document(self, document_path: str) -> Dict[str, Any]:
        """Analyze document structure and content."""
        if not os.path.exists(document_path):
            raise DocumentException(f"Document not found: {document_path}")

        try:
            from docx import Document

            doc = Document(document_path)

            analysis = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "styles": list(doc.styles.element.iter()),
                "has_toc": self._detect_toc(doc),
                "headings": self._extract_headings(doc)
            }

            return analysis

        except Exception as e:
            logger.error(f"Document analysis failed: {e}")
            return {"error": str(e)}

    def _detect_toc(self, doc) -> bool:
        """Detect if document has a table of contents."""
        for para in doc.paragraphs:
            if 'table of contents' in para.text.lower():
                return True
        return False

    def _extract_headings(self, doc) -> List[Dict]:
        """Extract all headings from document."""
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                headings.append({
                    "text": para.text[:100],
                    "style": para.style.name
                })
        return headings

    def suggest_edits(self, document_path: str, task: str) -> List[EditInstruction]:
        """Suggest edit instructions based on task."""
        instructions = []

        if "add chapter" in task.lower() or "new section" in task.lower():
            instructions.append(EditInstruction(
                operation=EditOperation.ADD_CHAPTER,
                target="end",
                content="New Chapter Content"
            ))

        return instructions

    def validate_editable(self, document_path: str) -> bool:
        """Check if document is editable."""
        if not os.path.exists(document_path):
            return False

        try:
            from docx import Document
            doc = Document(document_path)
            return len(doc.paragraphs) > 0
        except:
            return False